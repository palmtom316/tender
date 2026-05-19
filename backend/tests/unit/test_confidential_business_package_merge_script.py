from __future__ import annotations

from docx import Document

import scripts.merge_confidential_business_docx_package as merger


def _write_docx(path, *paragraphs: str) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


def test_merge_confidential_business_package_sorts_and_normalizes_without_content_in_evidence(tmp_path):
    source_dir = tmp_path / "package"
    source_dir.mkdir()
    _write_docx(source_dir / "8.1.1.资产负债表2022.docx", "8.1.1.资产负债表2022", "敏感正文 13800000000")
    _write_docx(source_dir / "8.近三年财务状况.docx", "8.近三年财务状况", "父章正文")
    _write_docx(source_dir / "9.9.联合体协议书.docx", "9.9.联合体协议书", "联合体正文")
    _write_docx(source_dir / "8.1.2022年财务会计报表.docx", "8.1.2022年财务会计报表", "年度正文")

    output = tmp_path / "merged.docx"
    evidence = merger.merge_package(source_dir, output)

    merged = Document(str(output))
    merged_text = "\n".join(paragraph.text for paragraph in merged.paragraphs)
    assert evidence["merge_order_codes"] == ["8", "8.1", "8.1.1", "9"]
    assert [item["normalized_code"] for item in evidence["items"]] == ["8", "8.1", "8.1.1", "9"]
    assert "9. 联合体协议书" in merged_text
    assert "9.9.联合体协议书" not in merged_text
    assert merged_text.index("8. 近三年财务状况") < merged_text.index("8.1. 2022年财务会计报表")
    assert merged_text.index("8.1. 2022年财务会计报表") < merged_text.index("8.1.1. 资产负债表2022")

    serialized = merger.to_json_text(evidence)
    assert "敏感正文" not in serialized
    assert "13800000000" not in serialized
    assert evidence["output_sha256"]
    assert evidence["output_size_bytes"] > 0
