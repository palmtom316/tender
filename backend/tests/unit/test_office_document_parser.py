from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from tender_backend.services.office_document_parser import OfficeParseError, parse_docx, parse_office_file


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "投标文件.docx"
    document = Document()
    document.add_heading("一、项目概况", level=1)
    document.add_paragraph("本项目为居民供电设施改造。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "人员"
    table.cell(0, 1).text = "证书"
    table.cell(1, 0).text = "项目经理"
    table.cell(1, 1).text = "建造师"
    document.save(path)

    chunks = parse_docx(path, source_file="投标文件.docx")

    assert any(chunk["chunk_type"] == "heading" and chunk["text"] == "一、项目概况" for chunk in chunks)
    assert any(chunk["chunk_type"] == "paragraph" and "居民供电设施改造" in chunk["text"] for chunk in chunks)
    table_chunks = [chunk for chunk in chunks if chunk["chunk_type"] == "table"]
    assert table_chunks
    assert table_chunks[0]["table_json"]["rows"][1] == ["项目经理", "建造师"]


def test_parse_legacy_doc_reports_missing_libreoffice(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    path = tmp_path / "投标文件.doc"
    path.write_bytes(b"legacy")
    monkeypatch.setattr("shutil.which", lambda _name: None)

    with pytest.raises(OfficeParseError, match="LibreOffice"):
        parse_office_file(path, source_file="投标文件.doc")
