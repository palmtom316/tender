from __future__ import annotations

import signal
from datetime import UTC, datetime
from pathlib import Path
from uuid import uuid4

from docx import Document

from tender_backend.db.repositories.chart_asset_repo import ChartAssetRow
from tender_backend.services.chart_service.captions import FigureNumbering
from tender_backend.services.chart_service.png_converter import svg_to_png
from tender_backend.services.chart_service.renderers import render_chart_spec
from tender_backend.services.chart_service.specs import parse_chart_spec
from tender_backend.services.export_service import chart_asset_injector
from tender_backend.services.export_service.chart_asset_injector import ChartAssetInjector


def test_native_risk_matrix_renders_svg_and_png(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("CHART_VEGA_ENGINE_ENABLED", "false")
    from tender_backend.core import config as _config

    _config.get_settings.cache_clear()
    spec = parse_chart_spec(
        {
            "chart_type": "risk_matrix",
            "title": "项目风险矩阵",
            "rows": ["低影响", "高影响"],
            "columns": ["低概率", "高概率"],
            "cells": [{"row": "高影响", "column": "高概率", "items": ["工期延误"], "level": "high"}],
        }
    )

    rendered = render_chart_spec(spec)
    png = svg_to_png(rendered.svg, tmp_path / "risk.png")

    assert "项目风险矩阵" in rendered.svg
    assert rendered.engine == "native_svg"
    assert png.is_file()
    assert png.stat().st_size > 0
    _config.get_settings.cache_clear()


def test_risk_matrix_uses_vl_convert_when_flag_enabled(tmp_path: Path, monkeypatch) -> None:
    import vl_convert  # noqa: F401

    monkeypatch.setenv("CHART_VEGA_ENGINE_ENABLED", "true")
    from tender_backend.core import config as _config

    _config.get_settings.cache_clear()
    spec = parse_chart_spec(
        {
            "chart_type": "risk_matrix",
            "title": "项目风险矩阵",
            "rows": ["低影响", "高影响"],
            "columns": ["低概率", "高概率"],
            "cells": [{"row": "高影响", "column": "高概率", "items": ["工期延误"], "level": "high"}],
        }
    )

    rendered = render_chart_spec(spec)

    assert rendered.engine == "vl_convert"
    assert rendered.svg.lstrip().startswith("<svg")
    assert "项目风险矩阵" in rendered.svg
    _config.get_settings.cache_clear()


def test_flow_fallback_renders_declared_edges_and_labels() -> None:
    spec = parse_chart_spec(
        {
            "chart_type": "quality_system",
            "title": "质量管理体系图",
            "nodes": [
                {"id": "manager", "label": "项目经理"},
                {"id": "quality", "label": "质量负责人"},
                {"id": "team", "label": "施工班组"},
            ],
            "edges": [
                {"from": "manager", "to": "quality", "label": "监督"},
                {"from": "manager", "to": "team", "label": "部署"},
            ],
        }
    )

    rendered = render_chart_spec(spec)

    assert rendered.engine == "system_mermaid_fallback"
    assert "监督" in rendered.svg
    assert "部署" in rendered.svg
    assert "data-edge='manager-quality'" in rendered.svg
    assert "data-edge='manager-team'" in rendered.svg


def test_flow_fallback_handles_cycles_without_hanging() -> None:
    spec = parse_chart_spec(
        {
            "chart_type": "closure_flow",
            "title": "问题闭环流程图",
            "nodes": [
                {"id": "check", "label": "检查"},
                {"id": "rectify", "label": "整改"},
            ],
            "edges": [
                {"from": "check", "to": "rectify"},
                {"from": "rectify", "to": "check"},
            ],
        }
    )

    def _timeout(_signum, _frame):
        raise TimeoutError("flow layout did not terminate")

    previous = signal.signal(signal.SIGALRM, _timeout)
    signal.setitimer(signal.ITIMER_REAL, 1)
    try:
        rendered = render_chart_spec(spec)
    finally:
        signal.setitimer(signal.ITIMER_REAL, 0)
        signal.signal(signal.SIGALRM, previous)

    assert "data-edge='check-rectify'" in rendered.svg
    assert "data-edge='rectify-check'" in rendered.svg


def test_gantt_fallback_renders_ticks_dependencies_sections_and_critical_path() -> None:
    spec = parse_chart_spec(
        {
            "chart_type": "schedule_gantt",
            "title": "施工进度计划图",
            "tasks": [
                {
                    "id": "prepare",
                    "label": "施工准备",
                    "start": "2026-06-01",
                    "end": "2026-06-05",
                    "group": "准备阶段",
                },
                {
                    "id": "install",
                    "label": "设备安装",
                    "start": "2026-06-06",
                    "end": "2026-06-20",
                    "group": "实施阶段",
                    "is_critical": True,
                },
            ],
            "dependencies": [{"from": "prepare", "to": "install"}],
        }
    )

    rendered = render_chart_spec(spec)

    assert rendered.engine == "system_mermaid_fallback"
    assert "准备阶段" in rendered.svg
    assert "实施阶段" in rendered.svg
    assert "data-tick=" in rendered.svg
    assert "data-dependency='prepare-install'" in rendered.svg
    assert "stroke='#d92d20'" in rendered.svg


def test_table_chart_renderer_outputs_grid_cells() -> None:
    spec = parse_chart_spec(
        {
            "chart_type": "indicator_table",
            "title": "绿色施工指标表",
            "columns": ["指标", "来源", "记录"],
            "rows": [{"cells": ["节水", "招标文件", "台账"]}],
        }
    )

    rendered = render_chart_spec(spec)

    assert rendered.engine == "native_svg"
    assert "绿色施工指标表" in rendered.svg
    assert "节水" in rendered.svg
    assert "招标文件" in rendered.svg


def test_risk_matrix_summarizes_dense_cells() -> None:
    spec = parse_chart_spec(
        {
            "chart_type": "risk_matrix",
            "title": "风险矩阵",
            "rows": ["高影响"],
            "columns": ["高概率"],
            "cells": [
                {
                    "row": "高影响",
                    "column": "高概率",
                    "items": ["风险1", "风险2", "风险3", "风险4"],
                    "level": "critical",
                }
            ],
        }
    )

    rendered = render_chart_spec(spec)

    assert "4项·极高" in rendered.svg
    assert "风险4" not in rendered.svg


def test_figure_numbering_uses_chapter_prefix() -> None:
    numbering = FigureNumbering()

    assert numbering.next_number(chapter_code="8.1") == "图8.1-1"
    assert numbering.next_number(chapter_code="8.2") == "图8.2-1"
    assert numbering.next_number(chapter_code="10.1.3") == "图10.1-1"
    assert numbering.next_number(chapter_code="10.1.8") == "图10.1-2"
    assert numbering.next_number(chapter_code="10.3.5") == "图10.3-1"
    assert numbering.next_number(chapter_code="9") == "图9-1"
    assert numbering.next_number(chapter_code="9", explicit="图A-1") == "图A-1"


def test_chart_asset_injector_replaces_placeholder_with_image_and_caption(tmp_path: Path, monkeypatch) -> None:
    project_id = uuid4()
    png = tmp_path / "chart.png"
    svg_to_png(
        "<svg xmlns='http://www.w3.org/2000/svg' width='120' height='60'><rect width='120' height='60' fill='white'/><text x='10' y='30'>测试图</text></svg>",
        png,
    )
    asset = ChartAssetRow(
        id=uuid4(),
        project_id=project_id,
        outline_node_id=None,
        chart_type="construction_flow",
        title="施工流程图",
        spec_json={"chapter_code": "8.1"},
        rendered_svg=None,
        rendered_path=None,
        placeholder_key="construction_flow_main",
        mermaid_source="flowchart TB",
        rendered_png_path=str(png),
        status="approved",
        version=1,
        template_instance_id=None,
        template_revision_no=None,
        is_stale_by_template=False,
        stale_by_template_revision_no=None,
        stale_by_template_block_id=None,
        template_stale_reason=None,
        metadata_json={},
        created_at=datetime.now(UTC),
        updated_at=datetime.now(UTC),
    )

    class _Repo:
        def find_for_placeholder(self, _conn, *, project_id, key):
            assert key == "construction_flow_main"
            return [asset]

    monkeypatch.setattr(chart_asset_injector, "ChartAssetRepository", lambda: _Repo())

    document = Document()
    document.add_paragraph("{{chart:construction_flow_main}}")
    output = tmp_path / "out.docx"

    count = ChartAssetInjector(document, object(), project_id=project_id).inject_all()
    document.save(str(output))
    rendered = Document(str(output))
    text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

    assert count == 1
    assert "{{chart:" not in text
    assert "图8.1-1 施工流程图" in text
    assert len(rendered.inline_shapes) == 1


def test_chart_asset_injector_preserves_placeholder_when_asset_missing(monkeypatch) -> None:
    class _Repo:
        def find_for_placeholder(self, _conn, *, project_id, key):
            return []

    monkeypatch.setattr(chart_asset_injector, "ChartAssetRepository", lambda: _Repo())

    document = Document()
    document.add_paragraph("{{chart:construction_flow}}")

    count = ChartAssetInjector(document, object(), project_id=uuid4()).inject_all()
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert count == 0
    assert "{{chart:construction_flow}}" in text
