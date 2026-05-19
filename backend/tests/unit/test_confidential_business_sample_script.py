from __future__ import annotations

from docx import Document

import scripts.inspect_confidential_business_sample as sample_inspector


def test_inspect_confidential_sample_writes_no_sensitive_text(tmp_path):
    sample = tmp_path / "国网配网工程商务标1-24章.docx"
    document = Document()
    document.add_paragraph("一、商务偏差表")
    document.add_paragraph("北京测试电力有限公司 联系人 13800000000")
    document.add_paragraph("统一社会信用代码 91110000123456789X")
    document.add_paragraph("二、无违法失信行为的承诺函")
    document.save(sample)

    evidence = sample_inspector.inspect_sample(sample)

    serialized = sample_inspector.to_json_text(evidence)
    assert evidence["docx_path"] == str(sample)
    assert evidence["sha256"]
    assert evidence["size_bytes"] > 0
    assert evidence["heading_codes"] == ["1", "2"]
    assert evidence["sensitive_scan"]["mobile_phone"]["count"] == 1
    assert evidence["sensitive_scan"]["company_name"]["count"] == 1
    assert evidence["sensitive_scan"]["unified_social_credit_code"]["count"] == 1
    assert "北京测试电力有限公司" not in serialized
    assert "13800000000" not in serialized
    assert "91110000123456789X" not in serialized
