from pathlib import Path

from docx import Document

from tender_backend.services.export_service.format_checker import check_docx_format


def test_check_docx_format_reports_missing_table_borders(tmp_path: Path):
    path = tmp_path / "bad-table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "无边框表格"
    document.save(path)

    result = check_docx_format(path)

    assert result["format_passed"] is False
    assert result["format_status"] == "failed"
    assert any(issue["code"] == "table_missing_borders" for issue in result["issues"])


def test_check_docx_format_passes_basic_document_without_tables(tmp_path: Path):
    path = tmp_path / "basic.docx"
    document = Document()
    paragraph = document.add_paragraph("正文内容")
    paragraph.style = document.styles["Normal"]
    document.save(path)

    result = check_docx_format(path)

    assert result["format_passed"] is True
    assert result["format_status"] == "passed"
