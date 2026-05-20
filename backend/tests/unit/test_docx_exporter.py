from __future__ import annotations

import zipfile
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document

from tender_backend.services.export_service import docx_exporter
from tender_backend.services.export_service.docx_exporter import (
    EXPORT_MODE_MULTI_DOC_ZIP,
    EXPORT_MODE_MULTI_DOCX_ZIP,
    EXPORT_MODE_SINGLE_DOCX,
    _is_skeleton_only,
    _render_plain_docx,
    render_chapter_docx_zip,
    render_docx,
    render_export,
    render_volume_docx,
)


class _Cursor:
    def __init__(self, drafts: list[dict] | None = None, facts: list[dict] | None = None):
        self._drafts = drafts
        self._facts = facts or []
        self.result = []

    def __enter__(self):
        return self

    def __exit__(self, *_args):
        return False

    def execute(self, query, params=None):
        if "SELECT name FROM project" in query:
            self.result = [("测试项目",)]
        elif "FROM chapter_draft" in query:
            self.result = self._drafts if self._drafts is not None else [
                {
                    "chapter_code": "1.1",
                    "content_md": "# 1.1 资格响应\n\n## 响应内容\n- 已提供营业执照",
                    "chapter_title": "资格响应",
                    "volume_type": "qualification",
                    "sort_order": 1,
                }
            ]
        elif "FROM project_requirement" in query:
            self.result = []
        elif "SELECT fact_key" in query:
            self.result = self._facts
        else:
            self.result = []
        return self

    def fetchone(self):
        return self.result[0] if self.result else None

    def fetchall(self):
        return self.result


class _Conn:
    def __init__(self, drafts: list[dict] | None = None, facts: list[dict] | None = None):
        self._drafts = drafts
        self._facts = facts or []

    def cursor(self, *args, **kwargs):
        return _Cursor(self._drafts, self._facts)


def test_render_docx_without_template_creates_plain_word_file(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("TEMPLATE_DIR", str(tmp_path / "missing-templates"))
    output = tmp_path / "out.docx"

    path = render_docx(_Conn(), project_id=uuid4(), output_path=output)

    assert path == output
    document = Document(str(output))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "测试项目 投标文件" in text
    assert "已提供营业执照" in text


def test_render_volume_docx_injects_equipment_tables_for_business_and_technical(tmp_path: Path, monkeypatch) -> None:
    injected: list[tuple[object, str]] = []

    class _Injector:
        def __init__(self, document, _conn, *, project_id):
            injected.append((document, str(project_id)))

        def inject_all(self):
            injected.append(("inject", "done"))
            return 4

    monkeypatch.setattr(docx_exporter, "EquipmentTableInjector", _Injector)
    monkeypatch.setattr(docx_exporter, "PersonnelTableInjector", _Injector)

    business_output = tmp_path / "business.docx"
    technical_output = tmp_path / "technical.docx"
    qualification_output = tmp_path / "qualification.docx"
    project_id = uuid4()

    render_volume_docx(_Conn(_multi_chapter_drafts()), project_id=project_id, volume_type="business", output_path=business_output)
    render_volume_docx(_Conn(_multi_chapter_drafts()), project_id=project_id, volume_type="technical", output_path=technical_output)
    render_volume_docx(_Conn(_multi_chapter_drafts()), project_id=project_id, volume_type="qualification", output_path=qualification_output)

    business_doc = Document(str(business_output))
    technical_doc = Document(str(technical_output))
    qualification_doc = Document(str(qualification_output))

    business_text = "\n".join(paragraph.text for paragraph in business_doc.paragraphs)
    technical_text = "\n".join(paragraph.text for paragraph in technical_doc.paragraphs)
    qualification_text = "\n".join(paragraph.text for paragraph in qualification_doc.paragraphs)

    assert "主要施工设备表" in business_text
    assert "{{equipment_table:vehicle}}" in business_text
    assert "项目管理机构人员表" in business_text
    assert "{{personnel_table}}" in business_text
    assert "主要施工设备表" in technical_text
    assert "{{equipment_table:vehicle}}" in technical_text
    assert "项目管理机构人员表" in technical_text
    assert "{{personnel_table}}" in technical_text
    assert "主要施工设备表" not in qualification_text
    assert "{{equipment_table:vehicle}}" not in qualification_text
    assert "项目管理机构人员表" not in qualification_text
    assert "{{personnel_table}}" not in qualification_text
    assert injected.count(("inject", "done")) == 4


def test_render_business_volume_docx_uses_rendered_chapter_artifact(tmp_path: Path) -> None:
    artifact = tmp_path / "chapter-5.docx"
    artifact_doc = Document()
    artifact_doc.add_heading("5 基本情况", level=1)
    artifact_doc.add_paragraph("来自单章 DOCX artifact 的正文")
    artifact_doc.save(artifact)

    output = tmp_path / "business.docx"
    drafts = [
        {
            "chapter_code": "5",
            "content_md": "fallback content should not be used",
            "chapter_title": "基本情况",
            "volume_type": "business",
            "sort_order": 5,
            "rendered_docx_path": str(artifact),
            "rendered_artifact_json": {"template_item_id": "template-5"},
        },
    ]

    render_volume_docx(_Conn(drafts), project_id=uuid4(), volume_type="business", output_path=output)

    document = Document(str(output))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "来自单章 DOCX artifact 的正文" in text
    assert "fallback content should not be used" not in text


def test_render_volume_docx_adds_three_part_header_and_cross_page_seal(tmp_path: Path) -> None:
    output = tmp_path / "business.docx"
    conn = _Conn(
        _multi_chapter_drafts(),
        facts=[{"fact_key": "company_name", "fact_value": "某某电力工程有限公司"}],
    )

    render_volume_docx(conn, project_id=uuid4(), volume_type="business", output_path=output)

    document = Document(str(output))
    header_text = "\n".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
    )
    header_table_text = "\n".join(
        cell.text
        for section in document.sections
        for table in section.header.tables
        for row in table.rows
        for cell in row.cells
    )
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "投标人=某某电力工程有限公司" in header_table_text
    assert "测试项目" in header_table_text
    assert "商务标" in header_table_text
    assert "投标文件" not in header_text
    assert body_text.count("骑缝章：____________") == len(_multi_chapter_drafts())


def _multi_chapter_drafts() -> list[dict]:
    return [
        {
            "chapter_code": "1.1",
            "content_md": "# 资格响应\n\n- 已提供营业执照",
            "chapter_title": "资格响应",
            "volume_type": "qualification",
            "sort_order": 1,
            "rendered_docx_path": None,
            "rendered_artifact_json": {},
        },
        {
            "chapter_code": "2.1",
            "content_md": "# 商务方案\n\n- 报价说明",
            "chapter_title": "商务方案",
            "volume_type": "business",
            "sort_order": 2,
            "rendered_docx_path": None,
            "rendered_artifact_json": {},
        },
    ]


def test_render_chapter_docx_zip_packs_each_chapter(tmp_path: Path) -> None:
    output = tmp_path / "chapters.zip"
    drafts = _multi_chapter_drafts()

    path = render_chapter_docx_zip(_Conn(drafts), project_id=uuid4(), output_path=output)

    assert path == output
    with zipfile.ZipFile(path) as archive:
        names = sorted(archive.namelist())
    assert len(names) == 2
    assert all(name.endswith(".docx") for name in names)
    assert names[0].startswith("01_")
    assert names[1].startswith("02_")
    assert "资格响应" in names[0]
    assert "商务方案" in names[1]


def test_render_chapter_docx_zip_raises_when_no_drafts(tmp_path: Path) -> None:
    output = tmp_path / "empty.zip"

    with pytest.raises(ValueError):
        render_chapter_docx_zip(_Conn(drafts=[]), project_id=uuid4(), output_path=output)


def test_render_export_dispatches_modes(tmp_path: Path, monkeypatch) -> None:
    captured: dict[str, object] = {}

    def fake_single(conn, *, project_id, template_name=None, output_path=None):
        captured["single"] = output_path
        target = output_path or tmp_path / "single.docx"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("single", encoding="utf-8")
        return target

    def fake_multi_docx(conn, *, project_id, output_path=None):
        captured["multi_docx"] = output_path
        target = output_path or tmp_path / "multi.docx.zip"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("docx-zip", encoding="utf-8")
        return target

    def fake_multi_doc(conn, *, project_id, output_path=None):
        captured["multi_doc"] = output_path
        target = output_path or tmp_path / "multi.doc.zip"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("doc-zip", encoding="utf-8")
        return target

    monkeypatch.setattr(docx_exporter, "render_docx", fake_single)
    monkeypatch.setattr(docx_exporter, "render_chapter_docx_zip", fake_multi_docx)
    monkeypatch.setattr(docx_exporter, "render_chapter_doc_zip", fake_multi_doc)

    project_id = uuid4()
    out_single = render_export(
        _Conn(),
        project_id=project_id,
        mode=EXPORT_MODE_SINGLE_DOCX,
        output_path=tmp_path / "s.docx",
    )
    out_docx_zip = render_export(
        _Conn(),
        project_id=project_id,
        mode=EXPORT_MODE_MULTI_DOCX_ZIP,
        output_path=tmp_path / "d.zip",
    )
    out_doc_zip = render_export(
        _Conn(),
        project_id=project_id,
        mode=EXPORT_MODE_MULTI_DOC_ZIP,
        output_path=tmp_path / "doc.zip",
    )

    assert out_single.read_text(encoding="utf-8") == "single"
    assert out_docx_zip.read_text(encoding="utf-8") == "docx-zip"
    assert out_doc_zip.read_text(encoding="utf-8") == "doc-zip"
    assert set(captured.keys()) == {"single", "multi_docx", "multi_doc"}


def test_render_export_rejects_unknown_mode(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        render_export(_Conn(), project_id=uuid4(), mode="unknown", output_path=tmp_path / "x")


def test_render_chapter_doc_zip_uses_doc_converter(tmp_path: Path, monkeypatch) -> None:
    output = tmp_path / "doc-zip.zip"
    drafts = _multi_chapter_drafts()

    def fake_convert(docx_path: Path) -> Path:
        doc_path = docx_path.with_suffix(".doc")
        doc_path.write_bytes(b"DOC")
        return doc_path

    monkeypatch.setattr(docx_exporter, "convert_docx_to_doc", fake_convert)

    path = docx_exporter.render_chapter_doc_zip(
        _Conn(drafts), project_id=uuid4(), output_path=output
    )

    assert path == output
    with zipfile.ZipFile(path) as archive:
        names = sorted(archive.namelist())
    assert all(name.endswith(".doc") for name in names)
    assert len(names) == 2


def test_render_chapter_doc_zip_raises_when_libreoffice_missing(tmp_path: Path, monkeypatch) -> None:
    output = tmp_path / "doc-zip.zip"
    monkeypatch.setattr(docx_exporter, "convert_docx_to_doc", lambda _path: None)

    with pytest.raises(RuntimeError):
        docx_exporter.render_chapter_doc_zip(
            _Conn(_multi_chapter_drafts()), project_id=uuid4(), output_path=output
        )


def test_export_evidence_reports_residual_chart_placeholders(tmp_path: Path, monkeypatch) -> None:
    from tender_backend.services.export_service.docx_exporter import inspect_rendered_docx_evidence

    path = tmp_path / "residual.docx"
    doc = Document()
    doc.add_paragraph("正文 {{chart:risk_matrix}}")
    doc.save(path)
    monkeypatch.setattr(
        "tender_backend.services.export_service.docx_exporter.count_docx_pages",
        lambda path: {"status": "unchecked", "actual_pages": None, "method": "test"},
    )

    evidence = inspect_rendered_docx_evidence(path)

    assert evidence["residual_chart_placeholders"] == ["risk_matrix"]
    assert evidence["residual_chart_placeholder_count"] == 1
    assert evidence["page_count"]["method"] == "test"


def test_is_skeleton_only_detects_boilerplate_only_content() -> None:
    content = """
# 8 施工方案与技术措施

## 8.1 编制依据
- 由项目经理统筹...

## 8.2 总体方案
- 围绕招标要求设置目标分解...
- 执行依据包括...
- 建立预警清单...
"""

    assert _is_skeleton_only(content) is True


def test_is_skeleton_only_accepts_real_content() -> None:
    content = "# 8 施工方案与技术措施\n\n## 8.1 编制依据\n" + ("本方案结合项目实际开展施工组织部署。" * 80)

    assert _is_skeleton_only(content) is False


def test_render_plain_docx_rejects_skeleton_only_content(tmp_path: Path) -> None:
    drafts = [
        {
            "chapter_code": "8",
            "content_md": "# 8 施工方案与技术措施\n\n## 8.1 编制依据\n- 由项目经理统筹...\n## 8.2 总体方案\n- 围绕招标要求设置目标分解...",
            "chapter_title": "施工方案与技术措施",
            "volume_type": "technical",
            "sort_order": 1,
            "metadata_json": {},
        }
    ]

    with pytest.raises(ValueError, match="skeleton-only"):
        _render_plain_docx(_Conn(drafts), project_id=uuid4(), output_path=tmp_path / "out.docx")

from tender_backend.services.export_service.docx_exporter import _should_add_chapter_divider_page


def test_technical_chapter_0_does_not_get_divider_page() -> None:
    assert _should_add_chapter_divider_page("technical", "0") is False
    assert _should_add_chapter_divider_page("technical", "1") is True
    assert _should_add_chapter_divider_page("technical", "0.1") is False
