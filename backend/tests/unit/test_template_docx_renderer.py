from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4

from docx import Document

from tender_backend.services.template_service.docx_renderer import (
    _render_company_profile,
    _render_financial_statements,
    _render_people,
    _render_performances,
    _sanitize_filename,
    render_template_item_docx,
)


def test_sanitize_filename_removes_unsafe_characters() -> None:
    assert _sanitize_filename("5.1_基本情况表.docx") == "5.1__.docx"
    assert _sanitize_filename("  ") == "rendered"


def test_sanitize_filename_keeps_ascii_safe_parts() -> None:
    assert _sanitize_filename("6.1_people-summary.docx") == "6.1_people-summary.docx"


def test_render_people_adds_summary_and_resume_sections() -> None:
    doc = Document()
    _render_people(
        doc,
        {
            "people": [
                {
                    "full_name": "唐玮",
                    "role_name": "项目总经理",
                    "title": "工程师",
                    "specialty": "机电工程",
                    "years_experience": 11,
                    "resume_text": "负责过多个配网服务项目。",
                }
            ]
        },
    )

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "拟委任的主要人员汇总表" in text
    assert "主要人员简历表（项目总经理）" in text
    assert "负责过多个配网服务项目。" in text


def test_render_company_profile_accepts_mapped_alias_fields() -> None:
    doc = Document()
    _render_company_profile(
        doc,
        {
            "company": {
                "company_title": "REDACTED",
                "credit_code": "91500105MA5U123456",
                "address_text": "重庆市江北区",
                "contact_summary": "王莉莉 / 13800000000",
            }
        },
    )

    table = doc.tables[0]
    values = [cell.text for row in table.rows for cell in row.cells]
    assert "REDACTED" in values
    assert "91500105MA5U123456" in values
    assert "重庆市江北区" in values
    assert "王莉莉 / 13800000000" in values


def test_render_people_accepts_mapped_alias_fields() -> None:
    doc = Document()
    _render_people(
        doc,
        {
            "people": [
                {
                    "person_name": "唐玮",
                    "role_label": "项目总经理",
                    "title_label": "工程师",
                    "specialty_label": "机电工程",
                    "experience_years_text": "11",
                    "contact_summary": "13800000000 / tang@example.com",
                    "resume_summary": "负责过多个配网服务项目。",
                }
            ]
        },
    )

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "主要人员简历表（项目总经理）" in text
    assert "负责过多个配网服务项目。" in text


def test_render_performances_adds_table_and_detail_sections() -> None:
    doc = Document()
    _render_performances(
        doc,
        {
            "performances": [
                {
                    "project_name": "渝中片区营配低压业务外包服务",
                    "client_name": "REDACTED市区供电分公司",
                    "contract_amount": 944.39,
                    "currency": "CNY",
                    "service_scope": "低压运维与抢修",
                    "started_on": "2024-01-01",
                    "ended_on": "2025-12-31",
                    "project_status": "已开工，未完工",
                }
            ]
        },
    )

    text = "\n".join(p.text for p in doc.paragraphs)
    assert "近年完成的类似项目情况表" in text
    assert "渝中片区营配低压业务外包服务" in text
    assert "低压运维与抢修" in text


def test_render_performances_accepts_mapped_alias_fields() -> None:
    doc = Document()
    _render_performances(
        doc,
        {
            "performances": [
                {
                    "project_title": "渝中片区营配低压业务外包服务",
                    "client_title": "REDACTED市区供电分公司",
                    "contract_amount_text": "944.39",
                    "service_scope_text": "低压运维与抢修",
                    "started_on_text": "2024-01-01",
                    "ended_on_text": "2025-12-31",
                    "contact_summary": "刘工 / 13900000000",
                }
            ]
        },
    )

    table = doc.tables[0]
    assert table.rows[1].cells[0].text == "渝中片区营配低压业务外包服务"
    assert table.rows[1].cells[2].text == "944.39"


def test_render_financial_statements_builds_year_matrix() -> None:
    doc = Document()
    _render_financial_statements(
        doc,
        {
            "financial_statements": [
                {
                    "fiscal_year": 2023,
                    "statement_type": "annual_report",
                    "statement_data": {"营业收入": "9509.69万元", "净利润": "968.93万元"},
                },
                {
                    "fiscal_year": 2024,
                    "statement_type": "annual_report",
                    "statement_data": {"营业收入": "7928.06万元", "净利润": "763.36万元"},
                },
            ]
        },
    )

    assert doc.tables[0].rows[0].cells[1].text == "2023"
    assert doc.tables[0].rows[0].cells[2].text == "2024"
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "近年财务状况表" in body_text


def test_render_template_item_docx_renders_single_docx_template(tmp_path: Path, monkeypatch) -> None:
    package_id = uuid4()
    item_id = uuid4()
    template_path = tmp_path / "商务标完整模板.docx"
    template = Document()
    template.add_paragraph("投标人：{{ company.company_name }}")
    template.save(template_path)

    item = SimpleNamespace(
        id=item_id,
        package_id=package_id,
        item_name="商务标完整模板",
        filename=template_path.name,
        relative_path=template_path.name,
        render_mode="single_docx",
        item_type="document",
    )
    package = SimpleNamespace(id=package_id, source_root=str(tmp_path))

    class _Repo:
        def get_item_by_id(self, conn, *, item_id):
            return item

        def get_by_id(self, conn, *, package_id):
            return package

    monkeypatch.setattr(
        "tender_backend.services.template_service.docx_renderer.BidTemplatePackageRepository",
        lambda: _Repo(),
    )
    monkeypatch.setattr(
        "tender_backend.services.template_service.docx_renderer.build_item_render_context",
        lambda conn, *, item_id: {
            "ready": True,
            "missing_required_bindings": [],
            "context": {"company": {"company_name": "REDACTED"}},
        },
    )

    result = render_template_item_docx(None, item_id=item_id, output_dir=tmp_path / "out")

    rendered = Document(result["output_path"])
    assert "投标人：REDACTED" in "\n".join(p.text for p in rendered.paragraphs)
