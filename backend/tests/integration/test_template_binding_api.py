from __future__ import annotations

import os
from pathlib import Path
from uuid import UUID

import fitz  # PyMuPDF
import psycopg
import pytest
from docx import Document

from tender_backend.core.config import get_settings
from tender_backend.db.migrations import load_initial_schema_sql
from tender_backend.main import app
from tender_backend.test_support.asgi_client import SyncASGIClient


_AUTH_HEADERS = {"Authorization": "Bearer dev-token"}


def _db_url() -> str | None:
    return os.environ.get("DATABASE_URL")


def _apply_schema(conn: psycopg.Connection) -> None:
    conn.execute(load_initial_schema_sql())
    conn.execute("""
    CREATE TABLE IF NOT EXISTS template_package_category (
      code TEXT PRIMARY KEY,
      display_name TEXT NOT NULL,
      description TEXT,
      sort_order INT NOT NULL DEFAULT 0,
      enabled BOOLEAN NOT NULL DEFAULT TRUE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS library_company (
      id UUID PRIMARY KEY,
      company_key TEXT NOT NULL UNIQUE,
      company_name TEXT NOT NULL,
      company_type TEXT,
      enabled BOOLEAN NOT NULL DEFAULT TRUE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS bid_template_package (
      id UUID PRIMARY KEY,
      package_key TEXT NOT NULL UNIQUE,
      display_name TEXT NOT NULL,
      package_type VARCHAR(32) NOT NULL,
      category_code TEXT REFERENCES template_package_category(code),
      source_root TEXT NOT NULL,
      source_manifest JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS bid_template_item (
      id UUID PRIMARY KEY,
      package_id UUID NOT NULL REFERENCES bid_template_package(id) ON DELETE CASCADE,
      item_code TEXT NULL,
      item_name TEXT NOT NULL,
      filename TEXT NOT NULL,
      relative_path TEXT NOT NULL,
      source_kind VARCHAR(16) NOT NULL DEFAULT 'docx',
      item_type VARCHAR(32) NOT NULL DEFAULT 'chapter',
      render_mode VARCHAR(32) NOT NULL DEFAULT 'templated',
      is_required BOOLEAN NOT NULL DEFAULT TRUE,
      sort_order INT NOT NULL DEFAULT 0,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      UNIQUE (package_id, relative_path)
    );
    CREATE TABLE IF NOT EXISTS bid_template_binding_rule (
      id UUID PRIMARY KEY,
      template_item_id UUID NOT NULL REFERENCES bid_template_item(id) ON DELETE CASCADE,
      binding_name TEXT NOT NULL,
      source_type VARCHAR(64) NOT NULL,
      selection_mode VARCHAR(32) NOT NULL DEFAULT 'all',
      source_filters JSONB NOT NULL DEFAULT '{}'::jsonb,
      field_mappings JSONB NOT NULL DEFAULT '[]'::jsonb,
      field_mapping_mode VARCHAR(16) NOT NULL DEFAULT 'augment',
      output_key TEXT NOT NULL,
      required BOOLEAN NOT NULL DEFAULT TRUE,
      sort_order INT NOT NULL DEFAULT 0,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      UNIQUE (template_item_id, binding_name)
    );
    CREATE TABLE IF NOT EXISTS company_profile (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      company_name TEXT NOT NULL,
      company_code TEXT,
      unified_social_credit_code TEXT,
      registered_address TEXT,
      contact_name TEXT,
      contact_phone TEXT,
      contact_email TEXT,
      website TEXT,
      registered_capital TEXT,
      company_type TEXT,
      business_scope TEXT,
      profile_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS person_profile (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      full_name TEXT NOT NULL,
      gender TEXT,
      age INT,
      education TEXT,
      title TEXT,
      role_name TEXT,
      specialty TEXT,
      years_experience INT,
      phone TEXT,
      email TEXT,
      resume_text TEXT,
      profile_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS project_performance (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      project_name TEXT NOT NULL,
      client_name TEXT NOT NULL,
      contract_amount NUMERIC(14,2),
      currency TEXT NOT NULL DEFAULT 'CNY',
      started_on DATE,
      ended_on DATE,
      project_status TEXT,
      service_scope TEXT,
      peak_staffing INT,
      average_staffing INT,
      contact_name TEXT,
      contact_phone TEXT,
      evidence_summary TEXT,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS qualification_certificate (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      certificate_name TEXT NOT NULL,
      certificate_type TEXT,
      certificate_no TEXT,
      holder_name TEXT,
      grade TEXT,
      specialty TEXT,
      issued_by TEXT,
      valid_from DATE,
      valid_to DATE,
      status TEXT NOT NULL DEFAULT 'active',
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    CREATE TABLE IF NOT EXISTS financial_statement (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      fiscal_year INT NOT NULL,
      statement_type TEXT NOT NULL,
      statement_data JSONB NOT NULL DEFAULT '{}'::jsonb,
      source_note TEXT,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      UNIQUE (fiscal_year, statement_type)
    );
    CREATE TABLE IF NOT EXISTS evidence_asset (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      owner_type TEXT NOT NULL,
      owner_id UUID,
      asset_name TEXT NOT NULL,
      asset_domain VARCHAR(64) NOT NULL DEFAULT 'generic',
      asset_category VARCHAR(64) NOT NULL DEFAULT 'supporting_document',
      asset_type TEXT NOT NULL DEFAULT 'supporting_document',
      file_name TEXT NOT NULL,
      file_path TEXT NOT NULL,
      media_type TEXT,
      issuer_name TEXT,
      issued_on DATE,
      expires_on DATE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      sort_order INT NOT NULL DEFAULT 0,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    """)
    conn.execute("""
    ALTER TABLE bid_template_binding_rule
      ADD COLUMN IF NOT EXISTS field_mappings JSONB NOT NULL DEFAULT '[]'::jsonb;
    """)
    conn.execute("""
    ALTER TABLE bid_template_binding_rule
      ADD COLUMN IF NOT EXISTS field_mapping_mode VARCHAR(16) NOT NULL DEFAULT 'augment';
    """)
    conn.execute("""
    ALTER TABLE bid_template_package
      ADD COLUMN IF NOT EXISTS category_code TEXT REFERENCES template_package_category(code);
    ALTER TABLE company_profile
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE person_profile
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE project_performance
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE qualification_certificate
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE financial_statement
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE evidence_asset
      ADD COLUMN IF NOT EXISTS library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE;
    ALTER TABLE evidence_asset
      ADD COLUMN IF NOT EXISTS asset_domain VARCHAR(64) NOT NULL DEFAULT 'generic';
    ALTER TABLE evidence_asset
      ADD COLUMN IF NOT EXISTS asset_category VARCHAR(64) NOT NULL DEFAULT 'supporting_document';
    """)
    conn.commit()


def _reset_schema(conn: psycopg.Connection) -> None:
    conn.execute("DELETE FROM bid_template_binding_rule;")
    conn.execute("DELETE FROM bid_template_item;")
    conn.execute("DELETE FROM bid_template_package;")
    conn.execute("DELETE FROM template_package_category;")
    conn.execute("DELETE FROM evidence_asset;")
    conn.execute("DELETE FROM financial_statement;")
    conn.execute("DELETE FROM qualification_certificate;")
    conn.execute("DELETE FROM project_performance;")
    conn.execute("DELETE FROM person_profile;")
    conn.execute("DELETE FROM company_profile;")
    conn.execute("DELETE FROM library_company;")
    conn.commit()


@pytest.fixture(autouse=True)
def _clear_settings_cache() -> None:
    get_settings.cache_clear()


def test_binding_rule_and_context_preview_flow(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    db_url = _db_url()
    if not db_url:
        pytest.skip("DATABASE_URL not set; skipping integration test")

    import_root = tmp_path / "imports"
    import_root.mkdir()
    source_dir = import_root / "20258B商务文件"
    source_dir.mkdir()
    template = Document()
    template.add_paragraph("投标人：{{ company.company_name }}")
    template.add_paragraph("项目人员数：{{ people|length }}")
    template.add_paragraph("附件数：{{ assets|length }}")
    template.save(source_dir / "20258B商务文件.docx")
    upload_root = tmp_path / "uploads"
    upload_root.mkdir()
    monkeypatch.setenv("TEMPLATE_IMPORT_ROOTS", str(import_root))
    monkeypatch.setenv("EVIDENCE_UPLOAD_DIR", str(upload_root))
    get_settings.cache_clear()

    with psycopg.connect(db_url) as conn:
        _apply_schema(conn)
        _reset_schema(conn)

    client = SyncASGIClient(app)
    client.headers.update(_AUTH_HEADERS)
    try:
        imported = client.post("/api/template-packages/import", json={"source_dir": str(source_dir)})
        assert imported.status_code == 200
        package = imported.json()
        package_id = UUID(package["id"])
        document_item_id = UUID(package["items"][0]["id"])

        company = client.post(
            "/api/master-data/company-profiles",
            json={"company_name": "REDACTED", "contact_name": "王莉莉"},
        )
        assert company.status_code == 201
        person = client.post(
            "/api/master-data/people",
            json={"full_name": "唐玮", "role_name": "项目总经理"},
        )
        assert person.status_code == 201
        certificate = client.post(
            "/api/master-data/certificates",
            json={
                "certificate_name": "质量管理体系认证证书",
                "holder_name": "REDACTED",
                "certificate_no": "ISO-001",
            },
        )
        assert certificate.status_code == 201
        certificate_id = UUID(certificate.json()["id"])
        attachment_file = upload_root / "quality-cert.pdf"
        pdf = fitz.open()
        page = pdf.new_page(width=595, height=842)
        page.insert_text((72, 72), "Quality Certificate", fontsize=24)
        pdf.save(attachment_file)
        pdf.close()
        evidence_asset = client.post(
            "/api/master-data/evidence-assets",
            json={
                "owner_type": "qualification_certificate",
                "owner_id": str(certificate_id),
                "asset_name": "质量认证证书扫描件",
                "asset_type": "certificate_scan",
                "file_name": "quality-cert.pdf",
                "file_path": str(attachment_file),
            },
        )
        assert evidence_asset.status_code == 201

        binding1 = client.post(
            f"/api/template-items/{document_item_id}/bindings",
            json={
                "binding_name": "company_basic",
                "source_type": "company_profile",
                "selection_mode": "latest",
                "field_mappings": [
                    {"target_field": "company_title", "source_field": "company_name"},
                    {"target_field": "contact_summary", "source_fields": ["contact_name", "contact_phone"], "transform": "join", "join_with": " / "},
                ],
                "output_key": "company",
            },
        )
        assert binding1.status_code == 201

        reimported = client.post("/api/template-packages/import", json={"source_dir": str(source_dir)})
        assert reimported.status_code == 200
        reimported_body = reimported.json()
        assert UUID(reimported_body["id"]) == package_id
        assert UUID(reimported_body["items"][0]["id"]) == document_item_id

        preserved_bindings = client.get(f"/api/template-items/{document_item_id}/bindings")
        assert preserved_bindings.status_code == 200
        assert [row["binding_name"] for row in preserved_bindings.json()] == ["company_basic"]

        binding2 = client.post(
            f"/api/template-items/{document_item_id}/bindings",
            json={
                "binding_name": "team_people",
                "source_type": "person_profile",
                "selection_mode": "all",
                "output_key": "people",
                "source_filters": {"equals": {"role_name": "项目总经理"}},
            },
        )
        assert binding2.status_code == 201
        binding2_id = UUID(binding2.json()["id"])
        binding3 = client.post(
            f"/api/template-items/{document_item_id}/bindings",
            json={
                "binding_name": "certificate_assets",
                "source_type": "evidence_asset",
                "selection_mode": "all",
                "output_key": "assets",
                "source_filters": {
                    "equals": {"owner_type": "qualification_certificate"},
                    "record_ids": [evidence_asset.json()["id"]],
                },
            },
        )
        assert binding3.status_code == 201

        preview = client.get(f"/api/template-packages/{package_id}/context-preview")
        assert preview.status_code == 200
        body = preview.json()
        preview_bindings = {row["binding_name"]: row for row in body["items"][0]["bindings"]}
        assert preview_bindings["company_basic"]["data"]["company_name"] == "REDACTED"
        assert preview_bindings["company_basic"]["data"]["company_title"] == "REDACTED"
        assert preview_bindings["company_basic"]["data"]["contact_summary"] == "王莉莉"
        assert preview_bindings["team_people"]["matched_count"] == 1
        assert preview_bindings["certificate_assets"]["matched_count"] == 1

        item_render = client.get(f"/api/template-items/{document_item_id}/render-context")
        assert item_render.status_code == 200
        assert item_render.json()["ready"] is True
        assert item_render.json()["context"]["company"]["company_name"] == "REDACTED"
        assert item_render.json()["context"]["company"]["company_title"] == "REDACTED"
        assert item_render.json()["bindings"][0]["field_mapping_mode"] == "augment"

        suggestions = client.get(f"/api/template-items/{document_item_id}/field-mapping-suggestions")
        assert suggestions.status_code == 200
        assert any(
            mapping["target_field"] == "company_title"
            for mapping in suggestions.json()["suggestions"][0]["field_mappings"]
        )

        package_render = client.get(f"/api/template-packages/{package_id}/render-context")
        assert package_render.status_code == 200
        assert package_render.json()["ready_item_count"] == 1
        assert package_render.json()["total_item_count"] == 1

        preflight = client.get(f"/api/template-packages/{package_id}/render-preflight")
        assert preflight.status_code == 200
        preflight_body = preflight.json()
        assert preflight_body["ready"] is True
        assert preflight_body["ready_item_count"] == 1
        assert preflight_body["blocked_item_count"] == 0
        assert set(preflight_body["items"][0]["context_keys"]) == {"assets", "company", "people"}

        rendered = client.post(f"/api/template-items/{document_item_id}/render-docx")
        assert rendered.status_code == 200
        assert rendered.json()["ready"] is True
        assert rendered.json()["output_path"].endswith(".docx")
        rendered_doc = Document(rendered.json()["output_path"])
        rendered_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs)
        assert "投标人：REDACTED" in rendered_text
        assert "项目人员数：1" in rendered_text
        assert "附件数：1" in rendered_text

        bundle = client.post(
            f"/api/template-packages/{package_id}/render-bundle",
            json={"include_zip": False},
        )
        assert bundle.status_code == 200
        bundle_body = bundle.json()
        assert bundle_body["rendered_count"] == 1
        assert bundle_body["failed_count"] == 0
        assert Path(bundle_body["output_dir"]).exists()
        assert bundle_body["items"][0]["output_path"].endswith(".docx")
        assert Path(bundle_body["items"][0]["output_path"]).exists()
        assert bundle_body["zip_path"] is None

        zipped_bundle = client.post(
            f"/api/template-packages/{package_id}/render-bundle",
            json={"include_zip": True},
        )
        assert zipped_bundle.status_code == 200
        zipped_body = zipped_bundle.json()
        assert zipped_body["zip_path"] is not None
        assert zipped_body["zip_path"].endswith(".zip")
        assert Path(zipped_body["zip_path"]).exists()

        updated = client.put(
            f"/api/template-bindings/{binding2_id}",
            json={"selection_mode": "first"},
        )
        assert updated.status_code == 200
        assert updated.json()["selection_mode"] == "first"

        listed = client.get(f"/api/template-items/{document_item_id}/bindings")
        assert listed.status_code == 200
        assert len(listed.json()) == 2

        deleted = client.delete(f"/api/template-bindings/{binding2_id}")
        assert deleted.status_code == 200
        assert deleted.json()["deleted"] is True

        anon_client = SyncASGIClient(app)
        unauthorized = anon_client.get(f"/api/template-items/{document_item_id}/bindings")
        assert unauthorized.status_code == 401
    finally:
        client.close()
        with psycopg.connect(db_url) as conn:
            _reset_schema(conn)
