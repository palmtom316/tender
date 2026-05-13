from __future__ import annotations

import os
from pathlib import Path
from uuid import UUID

import psycopg
import pytest
from docx import Document

from tender_backend.core.config import get_settings
from tender_backend.db.migrations import load_initial_schema_sql
from tender_backend.main import app
from tender_backend.test_support.asgi_client import SyncASGIClient


_AUTH_HEADERS = {"Authorization": "Bearer dev-token"}


def _db_url() -> str | None:
    return os.environ.get("DATABASE_URL")


def _apply_template_schema(conn: psycopg.Connection) -> None:
    conn.execute(load_initial_schema_sql())
    conn.execute("""
    CREATE TABLE IF NOT EXISTS template_package_category (
      code TEXT PRIMARY KEY,
      display_name TEXT NOT NULL,
      description TEXT,
      sort_order INT NOT NULL DEFAULT 0,
      enabled BOOLEAN NOT NULL DEFAULT TRUE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS bid_template_package (
      id UUID PRIMARY KEY,
      package_key TEXT NOT NULL UNIQUE,
      display_name TEXT NOT NULL,
      package_type VARCHAR(32) NOT NULL,
      category_code TEXT REFERENCES template_package_category(code),
      source_root TEXT NOT NULL,
      source_manifest JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS bid_template_item (
      id UUID PRIMARY KEY,
      package_id UUID NOT NULL REFERENCES bid_template_package(id) ON DELETE CASCADE,
      item_code TEXT NULL,
      item_name TEXT NOT NULL,
      filename TEXT NOT NULL,
      relative_path TEXT NOT NULL,
      source_kind VARCHAR(16) NOT NULL DEFAULT 'docx',
      item_type VARCHAR(32) NOT NULL DEFAULT 'chapter',
      render_mode VARCHAR(32) NOT NULL DEFAULT 'templated',
      is_required BOOLEAN NOT NULL DEFAULT TRUE,
      sort_order INT NOT NULL DEFAULT 0,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      UNIQUE (package_id, relative_path)
    );
    """)
    conn.execute("""
    ALTER TABLE bid_template_package
      ADD COLUMN IF NOT EXISTS category_code TEXT REFERENCES template_package_category(code);
    """)
    conn.commit()


def _reset_template_tables(conn: psycopg.Connection) -> None:
    conn.execute("DELETE FROM bid_template_item;")
    conn.execute("DELETE FROM bid_template_package;")
    conn.execute("DELETE FROM template_package_category;")
    conn.commit()


@pytest.fixture(autouse=True)
def _clear_settings_cache() -> None:
    get_settings.cache_clear()


def test_import_and_list_template_packages(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    db_url = _db_url()
    if not db_url:
        pytest.skip("DATABASE_URL not set; skipping integration test")

    import_root = tmp_path / "imports"
    import_root.mkdir()
    source_dir = import_root / "20258B商务文件"
    source_dir.mkdir()
    template = Document()
    template.add_paragraph("投标人：{{ company.company_name }}")
    template.save(source_dir / "20258B商务文件.docx")
    monkeypatch.setenv("TEMPLATE_IMPORT_ROOTS", str(import_root))
    get_settings.cache_clear()

    with psycopg.connect(db_url) as conn:
        _apply_template_schema(conn)
        _reset_template_tables(conn)

    client = SyncASGIClient(app)
    client.headers.update(_AUTH_HEADERS)
    try:
        response = client.post(
            "/api/template-packages/import",
            json={"source_dir": str(source_dir)},
        )
        assert response.status_code == 200
        body = response.json()
        package_id = UUID(body["id"])
        assert body["package_type"] == "business"
        assert body["item_count"] == 1
        assert body["items"][0]["item_code"] is None
        assert body["items"][0]["item_type"] == "document"
        assert body["items"][0]["render_mode"] == "single_docx"

        listed = client.get("/api/template-packages")
        assert listed.status_code == 200
        assert any(UUID(row["id"]) == package_id for row in listed.json())

        detail = client.get(f"/api/template-packages/{package_id}")
        assert detail.status_code == 200
        assert detail.json()["items"][0]["item_name"] == "20258B商务文件"

        anon_client = SyncASGIClient(app)
        anon_response = anon_client.get("/api/template-packages")
        assert anon_response.status_code == 401
    finally:
        client.close()
        with psycopg.connect(db_url) as conn:
            _reset_template_tables(conn)


def test_upload_template_package_seeds_missing_category_rows(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    db_url = _db_url()
    if not db_url:
        pytest.skip("DATABASE_URL not set; skipping integration test")

    import_root = tmp_path / "imports"
    import_root.mkdir()
    monkeypatch.setenv("TEMPLATE_IMPORT_ROOTS", str(import_root))
    get_settings.cache_clear()

    template_path = tmp_path / "国网配网工程商务标.docx"
    template = Document()
    template.add_paragraph("投标人：{{ company.company_name }}")
    template.save(template_path)

    with psycopg.connect(db_url) as conn:
        _apply_template_schema(conn)
        _reset_template_tables(conn)

    client = SyncASGIClient(app)
    client.headers.update(_AUTH_HEADERS)
    try:
        with template_path.open("rb") as handle:
            response = client.post(
                "/api/template-packages/upload",
                data={
                    "project_type": "国网配网工程",
                    "template_kind": "business",
                    "display_name": "国网配网工程商务标",
                    "category_code": "sgcc_distribution",
                },
                files={
                    "file": (
                        template_path.name,
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )

        assert response.status_code == 201
        body = response.json()
        assert body["package_type"] == "business"
        assert body["category_code"] == "sgcc_distribution"
        assert body["item_count"] == 1

        listed = client.get("/api/template-packages?category_code=sgcc_distribution")
        assert listed.status_code == 200
        assert len(listed.json()) == 1
        assert listed.json()[0]["display_name"] == "国网配网工程商务标"

        categories = client.get("/api/template-package-categories")
        assert categories.status_code == 200
        assert any(row["code"] == "sgcc_distribution" for row in categories.json())
    finally:
        client.close()
        with psycopg.connect(db_url) as conn:
            _reset_template_tables(conn)


def test_project_business_template_preview_returns_chapters(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    db_url = _db_url()
    if not db_url:
        pytest.skip("DATABASE_URL not set; skipping integration test")

    import_root = tmp_path / "imports"
    import_root.mkdir()
    monkeypatch.setenv("TEMPLATE_IMPORT_ROOTS", str(import_root))
    get_settings.cache_clear()

    template_path = tmp_path / "国网配网工程商务标.docx"
    template = Document()
    template.add_paragraph("一、商务偏差表")
    template.add_paragraph("商务偏差表正文")
    template.add_page_break()
    template.add_paragraph("二、承诺函")
    template.add_paragraph("承诺函正文")
    template.save(template_path)

    with psycopg.connect(db_url) as conn:
        _apply_template_schema(conn)
        _reset_template_tables(conn)

    client = SyncASGIClient(app)
    client.headers.update(_AUTH_HEADERS)
    try:
        with template_path.open("rb") as handle:
            uploaded = client.post(
                "/api/template-packages/upload",
                data={
                    "project_type": "国网配网工程",
                    "template_kind": "business",
                    "display_name": "国网配网工程商务标",
                    "category_code": "sgcc_distribution",
                },
                files={
                    "file": (
                        template_path.name,
                        handle,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert uploaded.status_code == 201
        package_id = uploaded.json()["id"]

        project = client.post("/api/projects", json={"name": "预览项目", "industry": "power", "business_line": "sgcc_distribution"})
        assert project.status_code == 200
        project_id = project.json()["id"]

        confirmed = client.post(
            f"/api/projects/{project_id}/template-selection",
            json={"package_id": package_id},
        )
        assert confirmed.status_code == 200

        preview = client.get(f"/api/projects/{project_id}/business-template-preview")
        assert preview.status_code == 200
        assert [chapter["chapter_title"] for chapter in preview.json()["chapters"]] == ["商务偏差表", "承诺函"]
    finally:
        client.close()
        with psycopg.connect(db_url) as conn:
            _reset_template_tables(conn)
