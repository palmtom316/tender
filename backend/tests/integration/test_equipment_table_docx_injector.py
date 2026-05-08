from __future__ import annotations

import os
from pathlib import Path
from uuid import uuid4

import psycopg
import pytest
from docx import Document

from tender_backend.core.config import get_settings
from tender_backend.db.migrations import load_initial_schema_sql
from tender_backend.services.export_service.docx_exporter import render_docx


def _db_url() -> str | None:
    return os.environ.get("DATABASE_URL")


def _apply_docx_schema(conn: psycopg.Connection) -> None:
    conn.execute(load_initial_schema_sql())
    conn.execute("""
    CREATE TABLE IF NOT EXISTS library_company (
      id UUID PRIMARY KEY,
      company_key TEXT NOT NULL UNIQUE,
      company_name TEXT NOT NULL,
      company_type TEXT,
      enabled BOOLEAN NOT NULL DEFAULT TRUE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS evidence_asset (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      owner_type TEXT NOT NULL,
      owner_id UUID,
      asset_name TEXT NOT NULL,
      asset_domain VARCHAR(64) NOT NULL DEFAULT 'generic',
      asset_category VARCHAR(64) NOT NULL DEFAULT 'supporting_document',
      asset_type TEXT NOT NULL DEFAULT 'supporting_document',
      file_name TEXT NOT NULL,
      file_path TEXT NOT NULL,
      media_type TEXT,
      issuer_name TEXT,
      issued_on DATE,
      expires_on DATE,
      metadata_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      sort_order INT NOT NULL DEFAULT 0,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS company_asset (
      id UUID PRIMARY KEY,
      library_company_id UUID NOT NULL REFERENCES library_company(id) ON DELETE CASCADE,
      asset_type TEXT NOT NULL CHECK (asset_type IN ('vehicle','machine','tool','safety')),
      name TEXT NOT NULL,
      spec_model TEXT,
      serial_no TEXT,
      manufacturer TEXT,
      quantity NUMERIC(12,2) NOT NULL DEFAULT 1,
      unit TEXT NOT NULL,
      ownership TEXT NOT NULL CHECK (ownership IN ('self','leased','third_party')),
      acquired_at DATE,
      expires_at DATE,
      technical_condition TEXT,
      status TEXT NOT NULL DEFAULT 'active' CHECK (status IN ('active','maintenance','retired')),
      location TEXT,
      extras JSONB NOT NULL DEFAULT '{}'::jsonb,
      notes TEXT,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS project_equipment_selection (
      id UUID PRIMARY KEY,
      project_id UUID NOT NULL REFERENCES project(id) ON DELETE CASCADE,
      asset_id UUID NOT NULL REFERENCES company_asset(id) ON DELETE RESTRICT,
      asset_type TEXT NOT NULL,
      intended_role TEXT,
      snapshot_json JSONB,
      display_order INT NOT NULL DEFAULT 0,
      confirmed BOOLEAN NOT NULL DEFAULT FALSE,
      confirmed_at TIMESTAMPTZ,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS person_profile (
      id UUID PRIMARY KEY,
      library_company_id UUID REFERENCES library_company(id) ON DELETE CASCADE,
      full_name TEXT NOT NULL,
      gender TEXT,
      age INT,
      education TEXT,
      title TEXT,
      role_name TEXT,
      specialty TEXT,
      years_experience INT,
      phone TEXT,
      email TEXT,
      resume_text TEXT,
      profile_json JSONB NOT NULL DEFAULT '{}'::jsonb,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );

    CREATE TABLE IF NOT EXISTS project_personnel_selection (
      id UUID PRIMARY KEY,
      project_id UUID NOT NULL REFERENCES project(id) ON DELETE CASCADE,
      person_id UUID NOT NULL REFERENCES person_profile(id) ON DELETE RESTRICT,
      intended_role TEXT,
      snapshot_json JSONB,
      display_order INT NOT NULL DEFAULT 0,
      confirmed BOOLEAN NOT NULL DEFAULT FALSE,
      confirmed_at TIMESTAMPTZ,
      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
      updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
    );
    """)
    conn.commit()


@pytest.fixture(autouse=True)
def _clear_settings_cache() -> None:
    get_settings.cache_clear()


@pytest.mark.skipif(not _db_url(), reason="DATABASE_URL not set")
def test_render_docx_injects_equipment_tables(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    template_dir = tmp_path / "templates"
    export_root = tmp_path / "exports"
    template_dir.mkdir(parents=True)
    export_root.mkdir(parents=True)

    template = Document()
    template.add_paragraph("本投标项目拟投入主要施工设备如下：")
    template.add_paragraph("{{equipment_table:vehicle}}")
    template.save(template_dir / "equipment-template.docx")

    monkeypatch.setenv("TEMPLATE_DIR", str(template_dir))
    monkeypatch.setenv("TENDER_EXPORT_ROOT", str(export_root))

    project_id = uuid4()
    asset_id = uuid4()
    library_company_id = uuid4()

    with psycopg.connect(_db_url()) as conn:
        _apply_docx_schema(conn)
        conn.execute(
            "INSERT INTO library_company (id, company_key, company_name) VALUES (%s, %s, %s)",
            (library_company_id, f"REDACTED-{library_company_id.hex[:8]}", "REDACTED"),
        )
        conn.execute("INSERT INTO project (id, name) VALUES (%s, %s)", (project_id, "测试项目"))
        conn.execute(
            """
            INSERT INTO company_asset (
              id, library_company_id, asset_type, name, spec_model, serial_no, manufacturer, quantity, unit, ownership, extras
            ) VALUES (%s, %s, 'vehicle', '斗臂车', 'DFL5160', '渝A12345', '东风', 1, '辆', 'self', '{"vehicle_type":"aerial_bucket"}'::jsonb)
            """,
            (asset_id, library_company_id),
        )
        conn.execute(
            """
            INSERT INTO project_equipment_selection (
              id, project_id, asset_id, asset_type, intended_role, snapshot_json, confirmed, confirmed_at
            ) VALUES (%s, %s, %s, 'vehicle', '配电主线', %s::jsonb, TRUE, now())
            """,
            (
                uuid4(),
                project_id,
                asset_id,
                '{"name":"斗臂车","spec_model":"DFL5160","serial_no":"渝A12345","manufacturer":"东风","quantity":"1","unit":"辆","ownership":"self","technical_condition":"良好","extras":{"vehicle_type":"aerial_bucket"}}',
            ),
        )
        conn.commit()

        output = render_docx(conn, project_id=project_id, template_name="equipment-template.docx")

    rendered = Document(str(output))
    paragraphs = "\n".join(p.text for p in rendered.paragraphs)
    assert "{{equipment_table:" not in paragraphs
    assert len(rendered.tables) >= 1
    assert rendered.tables[0].rows[0].cells[1].text == "设备名称"
    assert rendered.tables[0].rows[1].cells[1].text == "斗臂车"


@pytest.mark.skipif(not _db_url(), reason="DATABASE_URL not set")
def test_render_docx_injects_personnel_table(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    template_dir = tmp_path / "templates"
    export_root = tmp_path / "exports"
    template_dir.mkdir(parents=True)
    export_root.mkdir(parents=True)

    template = Document()
    template.add_paragraph("本项目管理机构人员如下：")
    template.add_paragraph("{{personnel_table}}")
    template.save(template_dir / "personnel-template.docx")

    monkeypatch.setenv("TEMPLATE_DIR", str(template_dir))
    monkeypatch.setenv("TENDER_EXPORT_ROOT", str(export_root))

    project_id = uuid4()
    person_id = uuid4()
    library_company_id = uuid4()

    with psycopg.connect(_db_url()) as conn:
        _apply_docx_schema(conn)
        conn.execute(
            "INSERT INTO library_company (id, company_key, company_name) VALUES (%s, %s, %s)",
            (library_company_id, f"cq-personnel-{library_company_id.hex[:8]}", "REDACTED"),
        )
        conn.execute("INSERT INTO project (id, name) VALUES (%s, %s)", (project_id, "测试项目"))
        conn.execute(
            """
            INSERT INTO person_profile (id, library_company_id, full_name, role_name, specialty, title, education, years_experience)
            VALUES (%s, %s, '张三', '项目经理', '电力工程', '高级工程师', '本科', 12)
            """,
            (person_id, library_company_id),
        )
        conn.execute(
            """
            INSERT INTO project_personnel_selection (
              id, project_id, person_id, intended_role, snapshot_json, confirmed, confirmed_at
            ) VALUES (%s, %s, %s, '项目负责人', %s::jsonb, TRUE, now())
            """,
            (
                uuid4(),
                project_id,
                person_id,
                '{"full_name":"张三","intended_role":"项目负责人","gender":"男","age":36,"education":"本科","title":"高级工程师","specialty":"电力工程","years_experience":12,"attachments":[{"asset_category":"practice_certificate","metadata_json":{"cert_no":"渝123"}}]}',
            ),
        )
        conn.commit()

        output = render_docx(conn, project_id=project_id, template_name="personnel-template.docx")

    rendered = Document(str(output))
    paragraphs = "\n".join(p.text for p in rendered.paragraphs)
    assert "{{personnel_table}}" not in paragraphs
    assert len(rendered.tables) >= 1
    assert rendered.tables[0].rows[0].cells[1].text == "姓名"
    assert rendered.tables[0].rows[1].cells[1].text == "张三"
    assert rendered.tables[0].rows[1].cells[2].text == "项目负责人"
