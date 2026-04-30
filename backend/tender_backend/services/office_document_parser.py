"""Office document parsing for tender and bid source files."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document


OFFICE_PARSER_NAME = "office-document-parser"
OFFICE_PARSER_VERSION = "0.1.0"


class OfficeParseError(RuntimeError):
    pass


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _table_to_rows(table: Any) -> list[list[str]]:
    return [[_cell_text(cell.text) for cell in row.cells] for row in table.rows]


def _first_non_empty(values: list[str]) -> str | None:
    for value in values:
        if value.strip():
            return value.strip()
    return None


def parse_docx(path: Path, *, source_file: str) -> list[dict[str, Any]]:
    document = Document(str(path))
    chunks: list[dict[str, Any]] = []
    order = 0
    current_title: str | None = None

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
        is_heading = style_name.lower().startswith("heading") or style_name.startswith("标题")
        if is_heading:
            current_title = text
        chunks.append(
            {
                "chunk_type": "heading" if is_heading else "paragraph",
                "source_file": source_file,
                "source_locator": f"paragraph:{index + 1}",
                "title": text if is_heading else current_title,
                "text": text,
                "paragraph_index": index + 1,
                "sort_order": order,
                "metadata_json": {"style": style_name},
            }
        )
        order += 1

    for table_index, table in enumerate(document.tables):
        rows = _table_to_rows(table)
        if not any(any(cell for cell in row) for row in rows):
            continue
        title = _first_non_empty(rows[0]) if rows else current_title
        chunks.append(
            {
                "chunk_type": "table",
                "source_file": source_file,
                "source_locator": f"table:{table_index + 1}",
                "title": title,
                "text": "\n".join("\t".join(row) for row in rows),
                "table_json": {"rows": rows},
                "sort_order": order,
                "metadata_json": {"table_index": table_index + 1},
            }
        )
        order += 1

    return chunks


def parse_xlsx(path: Path, *, source_file: str) -> list[dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise OfficeParseError("openpyxl is required to parse .xlsx files") from exc

    workbook = load_workbook(filename=path, read_only=False, data_only=True)
    chunks: list[dict[str, Any]] = []
    order = 0
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows():
            values = [_cell_text(cell.value) for cell in row]
            if any(values):
                rows.append(values)
        if not rows:
            continue
        title = sheet.title
        chunks.append(
            {
                "chunk_type": "table",
                "source_file": source_file,
                "source_locator": f"sheet:{sheet.title}",
                "title": title,
                "text": "\n".join("\t".join(row) for row in rows),
                "table_json": {"sheet_name": sheet.title, "rows": rows},
                "sheet_name": sheet.title,
                "row_start": 1,
                "row_end": sheet.max_row,
                "sort_order": order,
                "metadata_json": {
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "merged_cells": [str(cell_range) for cell_range in sheet.merged_cells.ranges],
                },
            }
        )
        order += 1
    workbook.close()
    return chunks


def _convert_with_libreoffice(path: Path, *, target_ext: str) -> Path:
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if binary is None:
        raise OfficeParseError("LibreOffice is required to parse legacy Office files (.doc/.xls)")

    with tempfile.TemporaryDirectory(prefix="tender-office-convert-") as tmp:
        output_dir = Path(tmp)
        result = subprocess.run(
            [
                binary,
                "--headless",
                "--convert-to",
                target_ext.lstrip("."),
                "--outdir",
                str(output_dir),
                str(path),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            raise OfficeParseError(f"LibreOffice conversion failed: {result.stderr.strip() or result.stdout.strip()}")
        converted = output_dir / f"{path.stem}.{target_ext.lstrip('.')}"
        if not converted.is_file():
            raise OfficeParseError(f"LibreOffice conversion did not produce {target_ext}")
        persisted = path.with_suffix(f".converted{target_ext}")
        persisted.write_bytes(converted.read_bytes())
        return persisted


def parse_office_file(path: Path, *, source_file: str) -> tuple[str, list[dict[str, Any]]]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "docx", parse_docx(path, source_file=source_file)
    if suffix == ".xlsx":
        return "xlsx", parse_xlsx(path, source_file=source_file)
    if suffix == ".doc":
        converted = _convert_with_libreoffice(path, target_ext=".docx")
        return "doc", parse_docx(converted, source_file=source_file)
    if suffix == ".xls":
        converted = _convert_with_libreoffice(path, target_ext=".xlsx")
        return "xls", parse_xlsx(converted, source_file=source_file)
    raise OfficeParseError(f"Unsupported Office file type: {suffix}")
