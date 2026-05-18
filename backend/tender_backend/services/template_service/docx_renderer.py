from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path
from uuid import UUID

from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from psycopg import Connection

from tender_backend.core.config import get_settings
from tender_backend.core.path_safety import ensure_path_within_root
from tender_backend.db.repositories.bid_template_package_repo import BidTemplatePackageRepository
from tender_backend.services.template_service.context_preview import build_item_render_context
from tender_backend.services.template_service.package_importer import _docx_heading_match
from tender_backend.services.tender_requirement_priority import load_tender_requirement_overrides


def _sanitize_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_") or "rendered"


def _add_key_value_table(doc: Document, pairs: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in pairs:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = "" if value is None else str(value)


def _coerce_records(value: object) -> list[dict]:
    if value is None:
        return []
    if isinstance(value, dict):
        return [value]
    if isinstance(value, list):
        return [item for item in value if isinstance(item, dict)]
    return []


def _format_period(started: object, ended: object) -> str:
    start = str(started or "").strip()
    end = str(ended or "").strip()
    if start and end:
        return f"{start} 至 {end}"
    return start or end


def _pick(record: dict, *keys: str) -> object:
    for key in keys:
        value = record.get(key)
        if value not in {None, ""}:
            return value
    return None


def _render_company_profile(doc: Document, context: dict) -> None:
    company = context.get("company") or {}
    doc.add_paragraph("应答人基本情况表")
    pairs = [
        ("单位名称", _pick(company, "company_name", "company_title", "unit_name")),
        ("统一社会信用代码", _pick(company, "unified_social_credit_code", "credit_code")),
        ("单位地址", _pick(company, "registered_address", "address_text")),
        ("联系人", _pick(company, "contact_name", "contact_person")),
        ("联系电话", _pick(company, "contact_phone", "phone", "contact_summary")),
        ("联系邮箱", company.get("contact_email")),
        ("网址", company.get("website")),
        ("注册资本", company.get("registered_capital")),
        ("单位性质", company.get("company_type")),
        ("经营范围", company.get("business_scope")),
    ]
    _add_key_value_table(doc, pairs)
    extra_profile = company.get("profile_json") or {}
    if extra_profile:
        doc.add_paragraph("补充信息")
        _add_key_value_table(doc, [(key, value) for key, value in extra_profile.items()])


def _render_people(doc: Document, context: dict) -> None:
    people = _coerce_records(context.get("people"))
    if not people:
        people = _coerce_records(context.get("person"))

    doc.add_paragraph("拟委任的主要人员汇总表")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "姓名"
    headers[1].text = "岗位"
    headers[2].text = "职称"
    headers[3].text = "专业"
    headers[4].text = "从业年限"
    headers[5].text = "联系方式"

    for person in people:
        row = table.add_row().cells
        row[0].text = str(_pick(person, "full_name", "person_name", "name") or "")
        row[1].text = str(_pick(person, "role_name", "role_label", "position_name") or "")
        row[2].text = str(_pick(person, "title", "title_label") or "")
        row[3].text = str(_pick(person, "specialty", "specialty_label", "profession") or "")
        experience = _pick(person, "years_experience", "experience_years_text", "experience_years")
        row[4].text = "" if experience is None else str(experience)
        row[5].text = str(_pick(person, "phone", "email", "contact_summary") or "")

    for person in people:
        doc.add_heading(
            f"主要人员简历表（{_pick(person, 'role_name', 'role_label', 'full_name', 'person_name') or '人员'}）",
            level=2,
        )
        _add_key_value_table(
            doc,
            [
                ("姓名", _pick(person, "full_name", "person_name", "name")),
                ("岗位", _pick(person, "role_name", "role_label", "position_name")),
                ("性别", person.get("gender")),
                ("年龄", person.get("age")),
                ("学历", person.get("education")),
                ("职称", _pick(person, "title", "title_label")),
                ("专业", _pick(person, "specialty", "specialty_label", "profession")),
                ("从业年限", _pick(person, "years_experience", "experience_years_text", "experience_years")),
                ("联系方式", _pick(person, "phone", "email", "contact_summary")),
            ],
        )
        resume_text = str(_pick(person, "resume_text", "resume_summary") or "").strip()
        if resume_text:
            doc.add_paragraph(resume_text)


def _render_performances(doc: Document, context: dict) -> None:
    performances = _coerce_records(context.get("performances"))
    if not performances:
        performances = _coerce_records(context.get("performance"))

    doc.add_paragraph("近年完成的类似项目情况表")
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "项目名称"
    headers[1].text = "业主单位"
    headers[2].text = "合同金额"
    headers[3].text = "服务范围"
    headers[4].text = "服务期间"
    headers[5].text = "人员配置"
    headers[6].text = "联系人"

    for performance in performances:
        row = table.add_row().cells
        row[0].text = str(_pick(performance, "project_name", "project_title") or "")
        row[1].text = str(_pick(performance, "client_name", "client_title") or "")
        amount_text = _pick(performance, "contract_amount_text")
        amount = performance.get("contract_amount")
        currency = performance.get("currency") or "CNY"
        row[2].text = str(amount_text) if amount_text not in {None, ""} else ("" if amount is None else f"{amount} {currency}")
        row[3].text = str(_pick(performance, "service_scope", "service_scope_text") or "")
        row[4].text = str(
            _pick(
                performance,
                "service_period_text",
                "period_text",
            )
            or _format_period(
                _pick(performance, "started_on_text", "started_on"),
                _pick(performance, "ended_on_text", "ended_on"),
            )
        )
        peak = performance.get("peak_staffing")
        avg = performance.get("average_staffing")
        staffing = []
        if peak is not None:
            staffing.append(f"高峰 {peak}")
        if avg is not None:
            staffing.append(f"平均 {avg}")
        row[5].text = str(_pick(performance, "staffing_summary") or " / ".join(staffing))
        row[6].text = str(_pick(performance, "contact_name", "contact_phone", "contact_summary") or "")

    for performance in performances:
        doc.add_heading(str(_pick(performance, "project_name", "project_title") or "项目说明"), level=2)
        doc.add_paragraph(str(_pick(performance, "service_scope", "service_scope_text") or ""))
        note_pairs = [
            ("项目状态", performance.get("project_status")),
            (
                "服务期间",
                _pick(performance, "service_period_text", "period_text")
                or _format_period(
                    _pick(performance, "started_on_text", "started_on"),
                    _pick(performance, "ended_on_text", "ended_on"),
                ),
            ),
            ("联系人", _pick(performance, "contact_name", "contact_summary")),
            ("联系电话", _pick(performance, "contact_phone", "contact_summary")),
            ("证明摘要", performance.get("evidence_summary")),
        ]
        _add_key_value_table(doc, note_pairs)


def _render_certificates(doc: Document, context: dict) -> None:
    certificates = _coerce_records(context.get("certificates"))
    if not certificates:
        certificates = _coerce_records(context.get("certificate"))

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "证书名称"
    headers[1].text = "证书编号"
    headers[2].text = "持有人"
    headers[3].text = "等级"
    headers[4].text = "有效期止"
    headers[5].text = "状态"

    for certificate in certificates:
        row = table.add_row().cells
        row[0].text = str(_pick(certificate, "certificate_name", "certificate_title") or "")
        row[1].text = str(_pick(certificate, "certificate_no", "certificate_code") or "")
        row[2].text = str(_pick(certificate, "holder_name", "holder_title") or "")
        row[3].text = str(_pick(certificate, "grade", "grade_label") or "")
        row[4].text = str(_pick(certificate, "valid_to", "valid_to_text") or "")
        row[5].text = str(_pick(certificate, "status", "status_label") or "")


def _render_financial_statements(doc: Document, context: dict) -> None:
    statements = _coerce_records(context.get("financial_statements"))
    if not statements:
        statements = _coerce_records(context.get("financial_statement"))

    doc.add_paragraph("近年财务状况表")
    if not statements:
        return

    years = [_pick(statement, "fiscal_year", "fiscal_year_label") for statement in statements]
    metric_names: list[str] = []
    metric_values_by_statement: list[dict] = []
    for statement in statements:
        data = statement.get("statement_data") or {}
        metric_values_by_statement.append(data)
        for key in data.keys():
            if key not in metric_names:
                metric_names.append(key)

    if metric_names:
        table = doc.add_table(rows=1, cols=len(statements) + 1)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "财务指标"
        for idx, year in enumerate(years, start=1):
            header[idx].text = str(year or "")
        for metric_name in metric_names:
            row = table.add_row().cells
            row[0].text = str(metric_name)
            for idx, data in enumerate(metric_values_by_statement, start=1):
                row[idx].text = "" if data.get(metric_name) is None else str(data.get(metric_name))

    for statement in statements:
        year = _pick(statement, "fiscal_year", "fiscal_year_label")
        statement_type = _pick(statement, "statement_type", "statement_title")
        doc.add_heading(f"{year} - {statement_type}", level=2)
        doc.add_paragraph(str(_pick(statement, "source_note", "source_note_text") or ""))


def _render_generic_context(doc: Document, context: dict) -> None:
    for key, value in context.items():
        doc.add_heading(str(key), level=2)
        if isinstance(value, list):
            for item in value:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph("" if value is None else str(value))


def _resolve_item_template_path(conn: Connection, item) -> Path:
    repo = BidTemplatePackageRepository()
    package = repo.get_by_id(conn, package_id=item.package_id)
    if package is None:
        raise LookupError("template package not found")

    root = Path(package.source_root)
    try:
        return ensure_path_within_root(root / item.relative_path, root, label="template item source path")
    except ValueError as exc:
        raise FileNotFoundError(str(exc)) from exc


def _split_docx_section_relative_path(relative_path: str, fallback_code: str | None = None) -> tuple[str, str]:
    docx_path, separator, anchor = relative_path.rpartition("#")
    chapter_code = (anchor or fallback_code or "").strip()
    if not separator or not docx_path.strip() or not chapter_code:
        raise ValueError("single_docx_section template item requires relative_path formatted as '<docx>#<chapter_code>'")
    return docx_path.strip(), chapter_code


def _code_parts(code: str) -> tuple[int, ...]:
    try:
        return tuple(int(part) for part in code.split("."))
    except ValueError:
        return ()


def _block_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def _extract_single_docx_section(source_path: Path, chapter_code: str, output_path: Path) -> None:
    shutil.copyfile(source_path, output_path)
    section_doc = Document(str(output_path))
    target_parts = _code_parts(chapter_code)
    if not target_parts:
        raise ValueError(f"invalid chapter code for single_docx_section: {chapter_code}")

    elements = [element for element in section_doc.element.body.iterchildren() if element.tag != qn("w:sectPr")]
    start_idx: int | None = None
    end_idx = len(elements)
    for idx, element in enumerate(elements):
        if element.tag == qn("w:sectPr"):
            continue
        heading = _docx_heading_match(_block_text(element)) if element.tag == qn("w:p") else None
        if heading is not None:
            current_code, _ = heading
            current_parts = _code_parts(current_code)
            if start_idx is not None and current_code != chapter_code and current_parts and len(current_parts) <= len(target_parts):
                end_idx = idx
                break
            if start_idx is None and current_code == chapter_code:
                start_idx = idx

    if start_idx is None:
        raise ValueError(f"chapter code not found in single DOCX template: {chapter_code}")
    if start_idx >= end_idx:
        raise ValueError(f"chapter section is empty in single DOCX template: {chapter_code}")
    keep_elements = set(elements[start_idx:end_idx])
    for element in elements:
        if element in keep_elements:
            continue
        section_doc.element.body.remove(element)
    section_doc.save(str(output_path))


def _render_single_docx_section_template(
    conn: Connection,
    *,
    item,
    context: dict,
    output_dir: Path | None,
    output_filename: str | None,
) -> Path:
    relative_docx_path, chapter_code = _split_docx_section_relative_path(item.relative_path, item.item_code)
    section_item = type(
        "_SectionTemplateItem",
        (),
        {"package_id": item.package_id, "relative_path": relative_docx_path},
    )()
    template_path = _resolve_item_template_path(conn, section_item)
    if not template_path.is_file():
        raise FileNotFoundError(f"template source file not found: {template_path}")

    root = output_dir or get_settings().template_render_root
    root.mkdir(parents=True, exist_ok=True)
    filename = output_filename or _sanitize_filename(f"{item.item_code or chapter_code}_{item.item_name}.docx")
    output_path = root / filename

    with tempfile.TemporaryDirectory(prefix="single-docx-section-") as temp_dir:
        section_template_path = Path(temp_dir) / "section.docx"
        _extract_single_docx_section(template_path, chapter_code, section_template_path)
        doc = DocxTemplate(str(section_template_path))
        doc.render(context)
        doc.save(str(output_path))
    return output_path


def _render_single_docx_template(
    conn: Connection,
    *,
    item,
    context: dict,
    output_dir: Path | None,
    output_filename: str | None,
) -> Path:
    template_path = _resolve_item_template_path(conn, item)
    if not template_path.is_file():
        raise FileNotFoundError(f"template source file not found: {template_path}")

    doc = DocxTemplate(str(template_path))
    doc.render(context)

    root = output_dir or get_settings().template_render_root
    root.mkdir(parents=True, exist_ok=True)
    filename = output_filename or _sanitize_filename(f"{item.item_name}.docx")
    output_path = root / filename
    doc.save(str(output_path))
    overrides = context.get("tender_requirement_priority")
    if overrides:
        rendered = Document(str(output_path))
        rendered.add_page_break()
        rendered.add_heading("招标文件解析要求优先响应", level=1)
        rendered.add_paragraph(str(overrides.get("description") or "招标文件解析要求优先于模板默认内容。"))
        for req in context.get("tender_content_requirements", []):
            rendered.add_paragraph(str(req.get("requirement_text") or req.get("source_text") or req.get("title") or ""))
        for req in context.get("tender_format_requirements", []):
            rendered.add_paragraph(str(req.get("requirement_text") or req.get("source_text") or req.get("title") or ""))
        rendered.save(str(output_path))
    return output_path


def render_template_item_docx(
    conn: Connection,
    *,
    item_id: UUID,
    output_dir: Path | None = None,
    output_filename: str | None = None,
    project_id: UUID | None = None,
) -> dict[str, object]:
    repo = BidTemplatePackageRepository()
    item = repo.get_item_by_id(conn, item_id=item_id)
    if item is None:
        raise LookupError("template item not found")

    if project_id is None:
        render_context = build_item_render_context(conn, item_id=item_id)
    else:
        render_context = build_item_render_context(conn, item_id=item_id, project_id=project_id)
    if not render_context["ready"]:
        missing = ", ".join(render_context["missing_required_bindings"])
        raise ValueError(f"template item is not ready for rendering: missing {missing}")

    context = render_context["context"]
    if item.render_mode == "single_docx_section":
        output_path = _render_single_docx_section_template(
            conn,
            item=item,
            context=context,
            output_dir=output_dir,
            output_filename=output_filename,
        )
        return {
            "item_id": str(item.id),
            "item_name": item.item_name,
            "filename": item.filename,
            "output_path": str(output_path),
            "ready": True,
            "context_keys": sorted(context.keys()),
        }

    if item.render_mode == "single_docx" or item.item_type == "document":
        output_path = _render_single_docx_template(
            conn,
            item=item,
            context=context,
            output_dir=output_dir,
            output_filename=output_filename,
        )
        return {
            "item_id": str(item.id),
            "item_name": item.item_name,
            "filename": item.filename,
            "output_path": str(output_path),
            "ready": True,
            "context_keys": sorted(context.keys()),
        }

    doc = Document()
    doc.add_heading(item.item_name, level=1)

    item_name = item.item_name
    item_code = item.item_code or ""
    if "基本情况表" in item_name:
        _render_company_profile(doc, context)
    elif "人员" in item_name or "项目团队" in item_name or item_code.startswith("6"):
        _render_people(doc, context)
    elif "业绩" in item_name or item_code.startswith("5"):
        _render_performances(doc, context)
    elif "证书" in item_name or "认证" in item_name:
        _render_certificates(doc, context)
    elif "财务" in item_name or item_code.startswith("8"):
        _render_financial_statements(doc, context)
    else:
        _render_generic_context(doc, context)

    root = output_dir or get_settings().template_render_root
    root.mkdir(parents=True, exist_ok=True)
    filename = output_filename or _sanitize_filename(f"{item_code or 'item'}_{item.item_name}.docx")
    output_path = root / filename
    doc.save(str(output_path))

    return {
        "item_id": str(item.id),
        "item_name": item.item_name,
        "filename": item.filename,
        "output_path": str(output_path),
        "ready": True,
        "context_keys": sorted(context.keys()),
    }


def render_project_template_blocks_for_preview(chapters: list[dict]) -> list[dict]:
    """Return DOCX-compatible preview markers for project template instance blocks.

    This intentionally keeps the boundary structured: final DOCX styling can consume
    these markers later without treating template editing as free-form source text.
    """
    rendered: list[dict] = []
    for chapter in chapters:
        blocks = []
        for block in chapter.get("blocks") or []:
            block_type = block.get("block_type")
            if block_type == "fixed_text":
                blocks.append({"kind": "paragraph", "text": block.get("content_text") or ""})
            elif block_type == "page_break":
                blocks.append({"kind": "page_break", "label": block.get("label")})
            elif block_type == "header_footer":
                blocks.append({"kind": "header_footer", "options": block.get("render_options_json") or {}})
            elif block_type == "seal_mark":
                blocks.append({"kind": "seal_mark", "label": block.get("label"), "metadata": block.get("metadata_json") or {}})
            else:
                blocks.append({"kind": "placeholder", "block_type": block_type, "label": block.get("label")})
        rendered.append({"chapter_code": chapter.get("chapter_code"), "chapter_title": chapter.get("chapter_title"), "blocks": blocks})
    return rendered
