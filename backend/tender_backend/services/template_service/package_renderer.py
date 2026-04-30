from __future__ import annotations

import io
import re
import shutil
from datetime import datetime
from pathlib import Path
from uuid import UUID
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.shared import Inches
from psycopg import Connection

from tender_backend.core.config import get_settings
from tender_backend.core.path_safety import ensure_path_within_root
from tender_backend.db.repositories.bid_template_package_repo import BidTemplatePackageRepository
from tender_backend.services.template_service.context_preview import build_item_render_context
from tender_backend.services.template_service.docx_renderer import render_template_item_docx
from tender_backend.services.vision_service.pdf_renderer import render_pdf_to_pages

_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}
_PDF_SUFFIXES = {".pdf"}
_LICENSE_MARKERS = ("营业执照", "法人证书", "登记证书", "开户许可证", "账户信息")
_CERTIFICATE_MARKERS = ("证书", "认证", "资质", "高新技术", "绿证", "许可证")
_AUTHORIZATION_MARKERS = ("授权", "委托", "身份证明", "承诺函", "说明")
_CONTRACT_MARKERS = ("合同", "保函", "保单", "汇款", "底单", "凭证")


def _to_cn_numeral(value: int) -> str:
    numerals = {
        1: "一",
        2: "二",
        3: "三",
        4: "四",
        5: "五",
        6: "六",
        7: "七",
        8: "八",
        9: "九",
        10: "十",
        11: "十一",
        12: "十二",
        13: "十三",
        14: "十四",
        15: "十五",
        16: "十六",
        17: "十七",
        18: "十八",
        19: "十九",
        20: "二十",
    }
    return numerals.get(value, str(value))


def _attachment_display_title(item_name: str, item_code: str | None) -> str:
    if item_code and "." in item_code:
        tail = item_code.split(".")[-1]
        if tail.isdigit():
            return f"（{_to_cn_numeral(int(tail))}）{item_name}"
    return item_name


def _attachment_template_kind(item_name: str, assets: list[dict[str, object]]) -> str:
    haystack = " ".join(
        [item_name, *(str(asset.get("asset_name") or "") for asset in assets), *(str(asset.get("asset_type") or "") for asset in assets)]
    )
    if any(marker in haystack for marker in _LICENSE_MARKERS):
        return "license"
    if any(marker in haystack for marker in _AUTHORIZATION_MARKERS):
        return "authorization"
    if any(marker in haystack for marker in _CONTRACT_MARKERS):
        return "contract"
    if any(marker in haystack for marker in _CERTIFICATE_MARKERS):
        return "certificate"
    return "generic"


def _sanitize_path_segment(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", value).strip(" ._")
    return cleaned or "bundle"


def _bundle_dir_name(display_name: str, package_key: str) -> str:
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    base = _sanitize_path_segment(display_name or package_key)
    return f"{base}_{stamp}"


def _safe_relative_path(relative_path: str) -> Path:
    path = Path(relative_path)
    if path.is_absolute() or ".." in path.parts:
        raise ValueError(f"unsafe relative path: {relative_path}")
    return path


def _create_zip_archive(bundle_dir: Path) -> Path:
    zip_path = bundle_dir.with_suffix(".zip")
    with ZipFile(zip_path, mode="w", compression=ZIP_DEFLATED) as archive:
        for path in sorted(bundle_dir.rglob("*")):
            if path.is_file():
                archive.write(path, arcname=str(Path(bundle_dir.name) / path.relative_to(bundle_dir)))
    return zip_path


def _collect_evidence_assets(value: object) -> list[dict[str, object]]:
    collected: list[dict[str, object]] = []
    seen: set[str] = set()

    def _walk(node: object) -> None:
        if isinstance(node, dict):
            if node.get("file_path") and node.get("file_name"):
                key = str(node.get("id") or f"{node.get('file_path')}::{node.get('file_name')}")
                if key not in seen:
                    seen.add(key)
                    collected.append(node)
            for child in node.values():
                _walk(child)
        elif isinstance(node, list):
            for child in node:
                _walk(child)

    _walk(value)
    return collected


def _validated_asset_source_path(asset: dict[str, object]) -> Path:
    settings = get_settings()
    try:
        source_path = ensure_path_within_root(
            str(asset["file_path"]),
            settings.evidence_upload_dir,
            label="attachment file path",
        )
    except ValueError as exc:
        raise FileNotFoundError(str(exc)) from exc
    if not source_path.is_file():
        raise FileNotFoundError(f"attachment file not found: {source_path}")
    return source_path


def _copy_attachment_file(asset: dict[str, object], attachment_dir: Path) -> tuple[Path, str]:
    source_path = _validated_asset_source_path(asset)

    file_name = _sanitize_path_segment(str(asset.get("file_name") or source_path.name))
    if "." not in file_name and source_path.suffix:
        file_name = f"{file_name}{source_path.suffix}"
    destination = attachment_dir / file_name
    suffix = 1
    while destination.exists():
        destination = attachment_dir / f"{destination.stem}_{suffix}{destination.suffix}"
        suffix += 1
    shutil.copy2(source_path, destination)
    return destination, source_path.name


def _is_image_asset(source_path: Path, asset: dict[str, object]) -> bool:
    media_type = str(asset.get("media_type") or "").lower()
    return source_path.suffix.lower() in _IMAGE_SUFFIXES or media_type.startswith("image/")


def _is_pdf_asset(source_path: Path, asset: dict[str, object]) -> bool:
    media_type = str(asset.get("media_type") or "").lower()
    return source_path.suffix.lower() in _PDF_SUFFIXES or media_type == "application/pdf"


def _max_inline_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    return max(1.0, float(section.page_width - section.left_margin - section.right_margin) / 914400)


def _embed_asset_preview(doc: Document, *, source_path: Path, asset: dict[str, object]) -> bool:
    width_inches = _max_inline_width_inches(doc)

    if _is_image_asset(source_path, asset):
        doc.add_picture(str(source_path), width=Inches(width_inches))
        return True

    if _is_pdf_asset(source_path, asset):
        pages = render_pdf_to_pages(str(source_path), dpi=150)
        for page in pages:
            doc.add_paragraph(f"第 {page.page_number} 页")
            doc.add_picture(io.BytesIO(page.png_bytes), width=Inches(width_inches))
        return True

    return False


def _get_asset_meta(asset: dict[str, object], key: str) -> str:
    metadata = asset.get("metadata_json")
    if isinstance(metadata, dict):
        value = metadata.get(key)
        if value is not None:
            return str(value)
    value = asset.get(key)
    return "" if value is None else str(value)


def _add_summary_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def _add_evidence_page(
    doc: Document,
    *,
    asset: dict[str, object],
    source_path: Path,
    relative_export: Path,
) -> bool:
    doc.add_page_break()
    doc.add_heading(str(asset.get("asset_name") or source_path.name), level=2)
    doc.add_paragraph(f"源文件: {source_path.name}")
    doc.add_paragraph(f"导出位置: {relative_export}")
    preview_embedded = _embed_asset_preview(doc, source_path=source_path, asset=asset)
    if not preview_embedded:
        doc.add_paragraph("该文件类型暂不支持直接嵌入预览，请打开同级附件文件。")
    return preview_embedded


def _render_license_template(
    doc: Document,
    *,
    item_name: str,
    item_code: str | None,
    evidence_rows: list[dict[str, object]],
) -> None:
    doc.add_heading(_attachment_display_title(item_name, item_code), level=1)
    doc.add_paragraph("以下附企业主体资格证照材料。")
    rows = [
        [
            str(row["asset"].get("asset_name") or ""),
            _get_asset_meta(row["asset"], "holder_name"),
            _get_asset_meta(row["asset"], "unified_social_credit_code"),
            _get_asset_meta(row["asset"], "registered_address"),
            str(row["asset"].get("expires_on") or ""),
        ]
        for row in evidence_rows
    ]
    _add_summary_table(doc, ["证照名称", "主体名称", "统一社会信用代码", "注册地址", "有效期止"], rows)


def _render_certificate_template(
    doc: Document,
    *,
    item_name: str,
    item_code: str | None,
    evidence_rows: list[dict[str, object]],
) -> None:
    doc.add_heading(_attachment_display_title(item_name, item_code), level=1)
    doc.add_paragraph("以下附资质、体系认证或专项证书材料。")
    rows = [
        [
            str(row["asset"].get("asset_name") or ""),
            _get_asset_meta(row["asset"], "certificate_no"),
            _get_asset_meta(row["asset"], "holder_name"),
            _get_asset_meta(row["asset"], "issuer_name"),
            str(row["asset"].get("expires_on") or ""),
        ]
        for row in evidence_rows
    ]
    _add_summary_table(doc, ["证书名称", "证书编号", "持有人", "发证机构", "有效期止"], rows)


def _render_authorization_template(
    doc: Document,
    *,
    item_name: str,
    item_code: str | None,
    evidence_rows: list[dict[str, object]],
) -> None:
    doc.add_heading(_attachment_display_title(item_name, item_code), level=1)
    doc.add_paragraph("以下附授权、委托或承诺类文件。")
    rows = [
        [
            str(row["asset"].get("asset_name") or ""),
            _get_asset_meta(row["asset"], "principal_name"),
            _get_asset_meta(row["asset"], "agent_name"),
            _get_asset_meta(row["asset"], "authorization_scope"),
            str(row["asset"].get("issued_on") or ""),
        ]
        for row in evidence_rows
    ]
    _add_summary_table(doc, ["文件名称", "授权主体", "被授权人", "授权事项", "签发日期"], rows)


def _render_contract_template(
    doc: Document,
    *,
    item_name: str,
    item_code: str | None,
    evidence_rows: list[dict[str, object]],
) -> None:
    doc.add_heading(_attachment_display_title(item_name, item_code), level=1)
    doc.add_paragraph("以下附合同、保函、保单或缴款凭证等材料。")
    rows = [
        [
            str(row["asset"].get("asset_name") or ""),
            _get_asset_meta(row["asset"], "contract_no"),
            _get_asset_meta(row["asset"], "party_a"),
            _get_asset_meta(row["asset"], "party_b"),
            str(row["asset"].get("issued_on") or ""),
        ]
        for row in evidence_rows
    ]
    _add_summary_table(doc, ["材料名称", "编号", "甲方/付款方", "乙方/收款方", "日期"], rows)


def _render_generic_attachment_template(
    doc: Document,
    *,
    item_name: str,
    item_code: str | None,
    evidence_rows: list[dict[str, object]],
) -> None:
    doc.add_heading(_attachment_display_title(item_name, item_code), level=1)
    doc.add_paragraph("附件材料清单")
    rows = [
        [
            str(row["asset"].get("asset_name") or ""),
            row["original_name"],
            str(row["asset"].get("owner_type") or ""),
            str(row["asset"].get("expires_on") or ""),
            str(row["relative_export"]),
        ]
        for row in evidence_rows
    ]
    _add_summary_table(doc, ["材料名称", "文件名", "归属类型", "有效期止", "导出位置"], rows)


def _render_attachment_manifest(
    *,
    item_name: str,
    item_code: str | None,
    output_docx_path: Path,
    attachment_dir: Path,
    assets: list[dict[str, object]],
) -> list[dict[str, object]]:
    attachment_dir.mkdir(parents=True, exist_ok=True)
    copied_assets: list[dict[str, object]] = []
    evidence_rows: list[dict[str, object]] = []

    for asset in assets:
        source_path = _validated_asset_source_path(asset)
        copied_path, original_name = _copy_attachment_file(asset, attachment_dir)
        relative_export = copied_path.relative_to(output_docx_path.parent)
        copied_assets.append(
            {
                "asset_id": str(asset.get("id") or ""),
                "asset_name": asset.get("asset_name"),
                "file_name": original_name,
                "output_path": str(copied_path),
                "relative_output_path": str(relative_export),
                "preview_embedded": False,
            }
        )
        evidence_rows.append(
            {
                "asset": asset,
                "source_path": source_path,
                "copied_path": copied_path,
                "original_name": original_name,
                "relative_export": relative_export,
            }
        )

    doc = Document()
    template_kind = _attachment_template_kind(item_name, assets)
    if template_kind == "license":
        _render_license_template(doc, item_name=item_name, item_code=item_code, evidence_rows=evidence_rows)
    elif template_kind == "certificate":
        _render_certificate_template(doc, item_name=item_name, item_code=item_code, evidence_rows=evidence_rows)
    elif template_kind == "authorization":
        _render_authorization_template(doc, item_name=item_name, item_code=item_code, evidence_rows=evidence_rows)
    elif template_kind == "contract":
        _render_contract_template(doc, item_name=item_name, item_code=item_code, evidence_rows=evidence_rows)
    else:
        _render_generic_attachment_template(doc, item_name=item_name, item_code=item_code, evidence_rows=evidence_rows)

    for idx, row in enumerate(evidence_rows):
        preview_embedded = _add_evidence_page(
            doc,
            asset=row["asset"],
            source_path=row["source_path"],
            relative_export=row["relative_export"],
        )
        copied_assets[idx]["preview_embedded"] = preview_embedded

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))
    return copied_assets


def _render_attachment_item(
    conn: Connection,
    *,
    item,
    bundle_dir: Path,
    project_id: UUID | None = None,
) -> dict[str, object]:
    if project_id is None:
        render_context = build_item_render_context(conn, item_id=item.id)
    else:
        render_context = build_item_render_context(conn, item_id=item.id, project_id=project_id)
    if not render_context["ready"]:
        missing = ", ".join(render_context["missing_required_bindings"])
        raise ValueError(f"template item is not ready for rendering: missing {missing}")

    assets = _collect_evidence_assets(render_context["context"])
    if not assets:
        raise ValueError("attachment item has no evidence assets to export")

    relative_path = _safe_relative_path(item.relative_path)
    output_docx_path = bundle_dir / relative_path
    attachment_dir = output_docx_path.parent / f"{output_docx_path.stem}_attachments"
    copied_assets = _render_attachment_manifest(
        item_name=item.item_name,
        item_code=item.item_code,
        output_docx_path=output_docx_path,
        attachment_dir=attachment_dir,
        assets=assets,
    )
    return {
        "output_path": str(output_docx_path),
        "copied_assets": copied_assets,
        "copied_asset_count": len(copied_assets),
        "embedded_preview_count": sum(1 for asset in copied_assets if asset.get("preview_embedded")),
        "context_keys": sorted(render_context["context"].keys()),
    }


def preflight_template_package_bundle(
    conn: Connection,
    *,
    package_id: UUID,
) -> dict[str, object]:
    repo = BidTemplatePackageRepository()
    package = repo.get_by_id(conn, package_id=package_id)
    if package is None:
        raise LookupError("template package not found")

    items = repo.list_items(conn, package_id=package_id)
    item_results: list[dict[str, object]] = []
    ready_item_count = 0
    blocked_item_count = 0
    total_issue_count = 0

    for item in items:
        issues: list[dict[str, object]] = []
        missing_required_bindings: list[str] = []
        asset_count = 0
        valid_asset_count = 0
        invalid_asset_count = 0
        context_keys: list[str] = []

        try:
            _safe_relative_path(item.relative_path)
        except ValueError as exc:
            issues.append({"code": "unsafe_relative_path", "message": str(exc)})

        try:
            render_context = build_item_render_context(conn, item_id=item.id)
            missing_required_bindings = list(render_context["missing_required_bindings"])
            context_keys = sorted(render_context["context"].keys())
            if missing_required_bindings:
                issues.append(
                    {
                        "code": "missing_required_bindings",
                        "message": f"missing required bindings: {', '.join(missing_required_bindings)}",
                        "bindings": missing_required_bindings,
                    }
                )

            if item.render_mode == "attachment" or item.item_type == "evidence":
                assets = _collect_evidence_assets(render_context["context"])
                asset_count = len(assets)
                if not assets:
                    issues.append(
                        {
                            "code": "missing_evidence_assets",
                            "message": "attachment item has no evidence assets to export",
                        }
                    )
                for asset in assets:
                    try:
                        _validated_asset_source_path(asset)
                        valid_asset_count += 1
                    except FileNotFoundError as exc:
                        invalid_asset_count += 1
                        issues.append(
                            {
                                "code": "invalid_evidence_asset",
                                "message": str(exc),
                                "asset_id": str(asset.get("id") or ""),
                                "asset_name": str(asset.get("asset_name") or ""),
                                "file_name": str(asset.get("file_name") or ""),
                            }
                        )
        except LookupError as exc:
            issues.append({"code": "lookup_error", "message": str(exc)})
        except ValueError as exc:
            issues.append({"code": "context_error", "message": str(exc)})

        ready = not issues
        if ready:
            ready_item_count += 1
        else:
            blocked_item_count += 1
            total_issue_count += len(issues)

        item_results.append(
            {
                "item_id": str(item.id),
                "item_name": item.item_name,
                "filename": item.filename,
                "relative_path": item.relative_path,
                "render_mode": item.render_mode,
                "item_type": item.item_type,
                "ready": ready,
                "issue_count": len(issues),
                "issues": issues,
                "missing_required_bindings": missing_required_bindings,
                "asset_count": asset_count,
                "valid_asset_count": valid_asset_count,
                "invalid_asset_count": invalid_asset_count,
                "context_keys": context_keys,
            }
        )

    return {
        "package_id": str(package.id),
        "package_key": package.package_key,
        "display_name": package.display_name,
        "package_type": package.package_type,
        "total_item_count": len(items),
        "ready_item_count": ready_item_count,
        "blocked_item_count": blocked_item_count,
        "issue_count": total_issue_count,
        "ready": blocked_item_count == 0,
        "items": item_results,
    }


def render_template_package_bundle(
    conn: Connection,
    *,
    package_id: UUID,
    include_zip: bool = False,
    output_root: Path | None = None,
    project_id: UUID | None = None,
) -> dict[str, object]:
    repo = BidTemplatePackageRepository()
    package = repo.get_by_id(conn, package_id=package_id)
    if package is None:
        raise LookupError("template package not found")

    items = repo.list_items(conn, package_id=package_id)
    bundle_root = output_root or get_settings().template_bundle_root
    bundle_dir = bundle_root / _bundle_dir_name(package.display_name, package.package_key)
    bundle_dir.mkdir(parents=True, exist_ok=True)

    item_results: list[dict[str, object]] = []
    rendered_count = 0
    failed_count = 0

    for item in items:
        try:
            relative_path = _safe_relative_path(item.relative_path)
            if item.render_mode == "attachment" or item.item_type == "evidence":
                rendered = _render_attachment_item(conn, item=item, bundle_dir=bundle_dir, project_id=project_id)
            else:
                item_output_dir = bundle_dir / relative_path.parent
                rendered = render_template_item_docx(
                    conn,
                    item_id=item.id,
                    output_dir=item_output_dir,
                    output_filename=relative_path.name,
                    **({"project_id": project_id} if project_id is not None else {}),
                )
            item_results.append(
                {
                    "item_id": str(item.id),
                    "item_name": item.item_name,
                    "filename": item.filename,
                    "relative_path": item.relative_path,
                    "render_mode": item.render_mode,
                    "status": "rendered",
                    "output_path": rendered["output_path"],
                    "copied_asset_count": rendered.get("copied_asset_count", 0),
                    "embedded_preview_count": rendered.get("embedded_preview_count", 0),
                }
            )
            rendered_count += 1
        except Exception as exc:
            item_results.append(
                {
                    "item_id": str(item.id),
                    "item_name": item.item_name,
                    "filename": item.filename,
                    "relative_path": item.relative_path,
                    "render_mode": item.render_mode,
                    "status": "failed",
                    "error": str(exc),
                }
            )
            failed_count += 1

    zip_path = str(_create_zip_archive(bundle_dir)) if include_zip else None
    return {
        "package_id": str(package.id),
        "package_key": package.package_key,
        "display_name": package.display_name,
        "package_type": package.package_type,
        "output_dir": str(bundle_dir),
        "total_item_count": len(items),
        "rendered_count": rendered_count,
        "failed_count": failed_count,
        "zip_path": zip_path,
        "items": item_results,
    }
