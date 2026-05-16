from __future__ import annotations

import re
from pathlib import Path
from uuid import UUID

import structlog
from docx.document import Document
from docx.shared import Inches
from psycopg import Connection

from tender_backend.db.repositories.chart_asset_repo import ChartAssetRepository, ChartAssetRow
from tender_backend.services.chart_service.captions import FigureNumbering, add_caption_after
from tender_backend.services.chart_service.png_converter import svg_to_png


_ANCHOR_RE = re.compile(r"\{\{chart:([A-Za-z][A-Za-z0-9_.:-]{0,127})\}\}")
logger = structlog.stdlib.get_logger(__name__)


class ChartAssetInjector:
    def __init__(
        self,
        document: Document,
        conn: Connection,
        *,
        project_id: UUID,
        formal: bool = True,
        max_width_inches: float = 6.3,
    ):
        self._document = document
        self._conn = conn
        self._project_id = project_id
        self._formal = formal
        self._max_width_inches = max_width_inches
        self._repo = ChartAssetRepository()
        self._numbering = FigureNumbering()

    def inject_all(self) -> int:
        count = 0
        for paragraph in list(self._document.paragraphs):
            match = _ANCHOR_RE.search(paragraph.text or "")
            if not match:
                continue
            key = match.group(1)
            try:
                asset = self._select_asset(key)
                image_path = self._image_path(asset)
            except ValueError as exc:
                logger.warning("chart_asset_unavailable", project_id=str(self._project_id), key=key, error=str(exc))
                continue
            paragraph.clear()
            paragraph.alignment = 1
            paragraph.add_run().add_picture(str(image_path), width=Inches(self._max_width_inches))
            figure_no = self._numbering.next_number(
                chapter_code=_chapter_code(asset),
                explicit=_figure_no(asset),
            )
            add_caption_after(paragraph, figure_no=figure_no, title=_caption_title(asset))
            count += 1
        return count

    def _select_asset(self, key: str) -> ChartAssetRow:
        candidates = self._repo.find_for_placeholder(self._conn, project_id=self._project_id, key=key)
        if not candidates:
            raise ValueError(f"chart asset not found for placeholder: {key}")
        exact = [row for row in candidates if row.placeholder_key == key]
        pool = exact or [row for row in candidates if row.chart_type == key]
        if exact and len(exact) > 1:
            pool = [row for row in exact if row.status == "approved"] or exact
        if not exact and len([row for row in pool if row.status == "approved"]) > 1:
            raise ValueError(f"ambiguous chart assets for placeholder: {key}")
        approved = [row for row in pool if row.status == "approved"]
        if self._formal and not approved:
            raise ValueError(f"chart asset is not approved for placeholder: {key}")
        selected = (approved or pool)[0]
        if not selected.rendered_png_path and not selected.rendered_svg:
            raise ValueError(f"chart asset has no rendered output for placeholder: {key}")
        return selected

    def _image_path(self, asset: ChartAssetRow) -> Path:
        if asset.rendered_png_path:
            path = Path(asset.rendered_png_path)
            if path.is_file():
                return path
        if not asset.rendered_svg:
            raise ValueError(f"chart asset has no rendered SVG: {asset.id}")
        path = Path("/tmp/tender-chart-assets") / str(self._project_id) / f"{asset.id}.png"
        return svg_to_png(asset.rendered_svg, path)


def _caption_title(asset: ChartAssetRow) -> str:
    spec_value = asset.spec_json.get("caption_title")
    if spec_value:
        return str(spec_value)
    metadata_value = asset.metadata_json.get("caption_title")
    if metadata_value:
        return str(metadata_value)
    return asset.title


def _figure_no(asset: ChartAssetRow) -> str | None:
    value = asset.spec_json.get("figure_no") or asset.metadata_json.get("figure_no")
    return str(value) if value else None


def _chapter_code(asset: ChartAssetRow) -> str | None:
    value = asset.spec_json.get("chapter_code") or asset.metadata_json.get("chapter_code")
    if value:
        return str(value)
    source = asset.metadata_json.get("source_context")
    if isinstance(source, dict) and source.get("chapter_code"):
        return str(source["chapter_code"])
    return None
