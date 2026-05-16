"""DOCX exporter — renders chapter drafts into a Word template using docxtpl."""

from __future__ import annotations

import os
import re
import zipfile
from pathlib import Path
from uuid import UUID

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from psycopg import Connection
from psycopg.rows import dict_row

from tender_backend.services.export_service.doc_converter import convert_docx_to_doc
from tender_backend.services.export_service.chart_asset_injector import ChartAssetInjector
from tender_backend.services.export_service.equipment_table_injector import EquipmentTableInjector
from tender_backend.services.export_service.personnel_table_injector import PersonnelTableInjector
from tender_backend.services.export_service.page_counter import count_docx_pages
from tender_backend.services.tender_requirement_priority import load_tender_requirement_overrides

logger = structlog.stdlib.get_logger(__name__)

CHART_PLACEHOLDER_RE = re.compile(r"\{\{chart:([A-Za-z][A-Za-z0-9_.:-]{0,127})\}\}")
_SKELETON_HEADING_RE = re.compile(r"^#{1,6}\s+.*$", re.MULTILINE)
_SKELETON_LIST_RE = re.compile(r"^\s*[-*]\s+", re.MULTILINE)
_SKELETON_PLACEHOLDER_RE = re.compile(r"\{\{[^{}]+\}\}")
_SKELETON_TABLE_RE = re.compile(r"^\s*\|.*\|\s*$", re.MULTILINE)

_SKELETON_BOILERPLATE_PHRASES: tuple[str, ...] = (
    "由项目经理统筹",
    "围绕招标要求设置",
    "执行依据包括",
    "建立预警清单",
    "按确认约束执行并形成记录",
    "责任可追溯",
    "资料可核验",
)

_EQUIPMENT_TABLE_SECTIONS: tuple[tuple[str, str], ...] = (
    ("车辆", "vehicle"),
    ("施工机械", "machine"),
    ("施工工器具", "tool"),
    ("安全设施设备及器具", "safety"),
)

TEMPLATE_DIR = Path(os.environ.get("TEMPLATE_DIR", "templates"))
EXPORT_ROOT = Path(os.environ.get("TENDER_EXPORT_ROOT", "/tmp/tender-exports"))

EXPORT_MODE_SINGLE_DOCX = "single_docx"
EXPORT_MODE_MULTI_DOCX_ZIP = "multi_docx_zip"
EXPORT_MODE_MULTI_DOC_ZIP = "multi_doc_zip"
EXPORT_MODES: tuple[str, ...] = (
    EXPORT_MODE_SINGLE_DOCX,
    EXPORT_MODE_MULTI_DOCX_ZIP,
    EXPORT_MODE_MULTI_DOC_ZIP,
)


def _add_markdown_content(document: Document, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)


def _add_deviation_table(document: Document, chapter_title: str, deviation_data: dict) -> None:
    """Add deviation table (商务偏差表/技术偏差表) to document."""
    # Add table title
    title = chapter_title if chapter_title else "偏差表"
    title_para = document.add_paragraph(title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = "宋体"

    has_deviation = deviation_data.get("has_deviation", False)
    items = deviation_data.get("items", [])

    # Create table with 5 columns
    # Add extra rows: 1 header + 1 default row + 1 "以下无正文" + 8 empty rows
    num_rows = 1 + 1 + 1 + 8 + len(items) if has_deviation else 1 + 1 + 1 + 8
    table = document.add_table(rows=num_rows, cols=5)
    table.style = "Table Grid"

    # Set column headers
    headers = ["序号", "采购文件条目号", "采购文件条款", "应答文件条款", "偏差说明"]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Center align and bold
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "宋体"
                run.font.size = Pt(10.5)

    # Add default row (no deviation)
    default_row = table.rows[1]
    default_row.cells[0].text = "1"
    default_row.cells[1].text = "采购文件全部条目号"
    default_row.cells[2].text = "采购文件全部条款"
    default_row.cells[3].text = "应答文件全部条款"
    default_row.cells[4].text = "无偏差"
    for cell in default_row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "宋体"
                run.font.size = Pt(10.5)

    # Add "以下无正文" row
    no_content_row = table.rows[2]
    no_content_row.cells[0].merge(no_content_row.cells[4])
    no_content_row.cells[0].text = "以下无正文"
    for paragraph in no_content_row.cells[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(10.5)

    # Add deviation items if any
    if has_deviation and items:
        for idx, item in enumerate(items):
            row_idx = 3 + idx
            row = table.rows[row_idx]
            row.cells[0].text = str(item.get("seq_number", idx + 2))
            row.cells[1].text = item.get("procurement_clause_number", "")
            row.cells[2].text = item.get("procurement_clause", "")
            row.cells[3].text = item.get("response_clause", "")
            row.cells[4].text = item.get("deviation_note", "")
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "宋体"
                        run.font.size = Pt(10.5)

    # Add declaration text below table
    declaration_text = (
        "应答人声明：针对本采购标的，除本表已列明偏差外，"
        "我们接受采购文件规定的其余全部技术条件，"
        "并承诺按照采购文件规定的技术条件提供对应产品和服务。"
    )
    declaration_para = document.add_paragraph(declaration_text)
    for run in declaration_para.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)


def _load_project_name(conn: Connection, project_id: UUID) -> str:
    with conn.cursor() as cur:
        row = cur.execute("SELECT name FROM project WHERE id = %s", (project_id,)).fetchone()
    return str(row[0]) if row and row[0] else str(project_id)


def _apply_basic_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        section.header.paragraphs[0].text = "投标文件"
        footer = section.footer.paragraphs[0]
        footer.text = "第 "
        run = footer.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        footer.add_run(" 页")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _convert_to_chinese_number(num: int) -> str:
    """Convert Arabic number to Chinese number (1->一, 2->二, etc.)"""
    chinese_nums = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                    "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
                    "二十一", "二十二", "二十三", "二十四", "二十五", "二十六", "二十七", "二十八", "二十九", "三十"]
    if 1 <= num < len(chinese_nums):
        return chinese_nums[num]
    return str(num)


def _add_chapter_divider_page(document: Document, chapter_code: str, chapter_title: str) -> None:
    """Add a centered chapter divider page with title and subtitle."""
    # Try to extract the numeric part from chapter_code
    try:
        # Handle codes like "1", "2", "3" etc.
        if "." not in chapter_code:
            chapter_num = int(chapter_code)
            chinese_num = _convert_to_chinese_number(chapter_num)
            formatted_title = f"{chinese_num}、{chapter_title}"
        else:
            # For sub-chapters like "1.1", "2.3", don't add divider page
            return
    except (ValueError, TypeError):
        # If we can't parse the number, use the original format
        formatted_title = f"{chapter_code}、{chapter_title}"

    # Add multiple empty paragraphs to push content to vertical center
    for _ in range(10):
        document.add_paragraph()

    # Add the main title (centered)
    title_para = document.add_paragraph(formatted_title)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = "宋体"

    # Add the subtitle (centered)
    subtitle_para = document.add_paragraph("（本页不编辑正文）")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = "宋体"

    # Add page break after the divider page
    document.add_page_break()


def _should_include_equipment_tables(volume_type: str | None) -> bool:
    return volume_type in (None, "business", "technical")


def _should_include_personnel_table(volume_type: str | None) -> bool:
    return volume_type in (None, "technical")


def _append_equipment_table_anchors(document: Document) -> None:
    document.add_heading("主要施工设备表", level=1)
    for title, asset_type in _EQUIPMENT_TABLE_SECTIONS:
        document.add_paragraph(title)
        document.add_paragraph(f"{{{{equipment_table:{asset_type}}}}}")


def _append_personnel_table_anchor(document: Document) -> None:
    document.add_heading("项目管理机构人员表", level=1)
    document.add_paragraph("{{personnel_table}}")


def _is_skeleton_only(content_md: str) -> bool:
    normalized = content_md or ""
    normalized = _SKELETON_HEADING_RE.sub("", normalized)
    normalized = _SKELETON_LIST_RE.sub("", normalized)
    normalized = _SKELETON_PLACEHOLDER_RE.sub("", normalized)
    normalized = _SKELETON_TABLE_RE.sub("", normalized)
    for phrase in _SKELETON_BOILERPLATE_PHRASES:
        normalized = normalized.replace(phrase, "")
    normalized = re.sub(r"\s+", "", normalized)
    return len(normalized) < 500


def _render_plain_docx(
    conn: Connection,
    *,
    project_id: UUID,
    output_path: Path,
    volume_type: str | None = None,
) -> Path:
    project_name = _load_project_name(conn, project_id)
    with conn.cursor(row_factory=dict_row) as cur:
        if volume_type:
            drafts = cur.execute(
                """
                SELECT cd.chapter_code, cd.volume_type, cd.content_md, bc.chapter_title, bc.sort_order, bc.metadata_json
                FROM chapter_draft cd
                LEFT JOIN bid_chapter bc
                  ON bc.project_id = cd.project_id
                 AND bc.volume_type = cd.volume_type
                 AND bc.chapter_code = cd.chapter_code
                WHERE cd.project_id = %s AND cd.volume_type = %s
                ORDER BY bc.sort_order NULLS LAST, cd.chapter_code
                """,
                (project_id, volume_type),
            ).fetchall()
        else:
            drafts = cur.execute(
                """
                SELECT cd.chapter_code, cd.volume_type, cd.content_md, bc.chapter_title, bc.sort_order, bc.metadata_json
                FROM chapter_draft cd
                LEFT JOIN bid_chapter bc
                  ON bc.project_id = cd.project_id
                 AND bc.volume_type = cd.volume_type
                 AND bc.chapter_code = cd.chapter_code
                WHERE cd.project_id = %s
                ORDER BY cd.volume_type, bc.sort_order NULLS LAST, cd.chapter_code
                """,
                (project_id,),
            ).fetchall()

    for draft in drafts:
        content_md = str(draft.get("content_md") or "")
        if str(draft.get("volume_type") or "") == "technical" and _is_skeleton_only(content_md):
            raise ValueError(
                f"chapter {draft.get('chapter_code') or '?'} content is skeleton-only, refusing to export. Run generation first."
            )

    document = Document()
    _apply_basic_style(document)
    title = f"{project_name} 投标文件"
    if volume_type:
        title += f"（{volume_type} 分册）"
    document.add_heading(title, level=0)
    document.add_paragraph("本文件由系统根据已确认招标约束、企业资料和章节草稿生成。")
    document.add_page_break()
    document.add_heading("目录", level=1)
    for draft in drafts:
        document.add_paragraph(f"{draft['chapter_code']} {draft.get('chapter_title') or ''}".strip())
    if _should_include_equipment_tables(volume_type):
        document.add_page_break()
        _append_equipment_table_anchors(document)
    if _should_include_personnel_table(volume_type):
        document.add_page_break()
        _append_personnel_table_anchor(document)
    document.add_heading("附件清单", level=1)
    document.add_paragraph("资质证书、人员证书、业绩证明等附件按招标文件要求随交付包一并提交。")
    document.add_heading("签章位置", level=1)
    document.add_paragraph("投标人盖章：____________")
    document.add_paragraph("法定代表人或授权代表签字：____________")
    document.add_page_break()

    # Track previous chapter to detect top-level chapters
    prev_chapter_code = None
    for index, draft in enumerate(drafts):
        chapter_code = draft.get("chapter_code", "")
        chapter_title = draft.get("chapter_title", "")
        metadata_json = draft.get("metadata_json") or {}

        # Add divider page for top-level chapters (no dot in chapter_code)
        if chapter_code and "." not in str(chapter_code):
            _add_chapter_divider_page(document, str(chapter_code), str(chapter_title))
        elif index:
            document.add_page_break()

        # Check if this chapter has deviation table data
        deviation_data = metadata_json.get("deviation_table")
        if deviation_data:
            _add_deviation_table(document, str(chapter_title), deviation_data)
        else:
            _add_markdown_content(document, draft["content_md"])

    if _should_include_equipment_tables(volume_type):
        EquipmentTableInjector(document, conn, project_id=project_id).inject_all()
    if _should_include_personnel_table(volume_type):
        PersonnelTableInjector(document, conn, project_id=project_id).inject_all()
    ChartAssetInjector(document, conn, project_id=project_id, formal=True).inject_all()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def render_docx(
    conn: Connection,
    *,
    project_id: UUID,
    template_name: str = "default_technical_bid.docx",
    output_path: Path | None = None,
) -> Path:
    """Render project drafts into a DOCX file using docxtpl.

    Template placeholders: {{SECTION_<chapter_code>}} or {{<chapter_code>}}
    """
    from docxtpl import DocxTemplate

    if output_path is None:
        EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
        output_path = EXPORT_ROOT / str(project_id) / f"bid-{project_id}.docx"

    template_path = TEMPLATE_DIR / template_name
    if not template_path.exists():
        return _render_plain_docx(conn, project_id=project_id, output_path=output_path)

    doc = DocxTemplate(str(template_path))

    # Load all chapter drafts
    with conn.cursor(row_factory=dict_row) as cur:
        drafts = cur.execute(
            "SELECT volume_type, chapter_code, content_md FROM chapter_draft WHERE project_id = %s",
            (project_id,),
        ).fetchall()

    # Build context: both SECTION_xxx and xxx placeholders
    context: dict[str, str] = {}
    for d in drafts:
        code = d["chapter_code"]
        volume = d.get("volume_type") or ""
        content = d["content_md"]
        context[f"SECTION_{code}"] = content
        context[code] = content
        if volume:
            context[f"SECTION_{volume}_{code}"] = content
            context[f"{volume}_{code}"] = content

    # Load project facts as additional context
    with conn.cursor(row_factory=dict_row) as cur:
        facts = cur.execute(
            "SELECT fact_key, fact_value FROM project_fact WHERE project_id = %s",
            (project_id,),
        ).fetchall()
    for f in facts:
        context[f["fact_key"]] = f["fact_value"]

    doc.render(context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    document = Document(str(output_path))
    EquipmentTableInjector(document, conn, project_id=project_id).inject_all()
    PersonnelTableInjector(document, conn, project_id=project_id).inject_all()
    ChartAssetInjector(document, conn, project_id=project_id, formal=True).inject_all()
    overrides = load_tender_requirement_overrides(conn, project_id=project_id)
    if overrides["content_requirements"] or overrides["format_requirements"]:
        document.add_page_break()
        document.add_heading("招标文件解析要求优先响应", level=1)
        document.add_paragraph("以下内容由招标文件解析结果生成，内容与格式要求优先于模板默认内容；如有冲突，按本节要求执行。")
        if overrides["content_requirements"]:
            document.add_heading("内容要求", level=2)
            for req in overrides["content_requirements"]:
                document.add_paragraph(str(req.get("requirement_text") or req.get("source_text") or req.get("title") or ""))
        if overrides["format_requirements"]:
            document.add_heading("格式要求", level=2)
            for req in overrides["format_requirements"]:
                document.add_paragraph(str(req.get("requirement_text") or req.get("source_text") or req.get("title") or ""))
    document.save(str(output_path))

    logger.info("docx_rendered", project_id=str(project_id), output=str(output_path))
    return output_path


def render_volume_docx(
    conn: Connection,
    *,
    project_id: UUID,
    volume_type: str,
    output_path: Path | None = None,
) -> Path:
    if output_path is None:
        output_path = EXPORT_ROOT / str(project_id) / f"bid-{volume_type}-{project_id}.docx"
    return _render_plain_docx(conn, project_id=project_id, volume_type=volume_type, output_path=output_path)


_FILENAME_SAFE_RE = re.compile(r"[\\/:*?\"<>|\r\n\t]+")


def _safe_filename_segment(value: str, fallback: str = "chapter") -> str:
    cleaned = _FILENAME_SAFE_RE.sub("_", value).strip(" ._-")
    return cleaned or fallback


def _load_chapter_drafts(conn: Connection, project_id: UUID) -> list[dict]:
    with conn.cursor(row_factory=dict_row) as cur:
        rows = cur.execute(
            """
            SELECT cd.chapter_code, cd.volume_type, cd.content_md, bc.chapter_title, bc.sort_order
            FROM chapter_draft cd
            LEFT JOIN bid_chapter bc
              ON bc.project_id = cd.project_id
             AND bc.volume_type = cd.volume_type
             AND bc.chapter_code = cd.chapter_code
            WHERE cd.project_id = %s
            ORDER BY cd.volume_type, bc.sort_order NULLS LAST, cd.chapter_code
            """,
            (project_id,),
        ).fetchall()
    return [dict(row) for row in rows]


def _render_chapter_docx(
    *,
    project_name: str,
    draft: dict,
    output_path: Path,
) -> Path:
    document = Document()
    _apply_basic_style(document)
    chapter_code = str(draft.get("chapter_code") or "").strip()
    chapter_title = str(draft.get("chapter_title") or "").strip()
    heading = " ".join(part for part in (chapter_code, chapter_title) if part) or "章节"
    document.add_heading(f"{project_name} 投标文件", level=0)
    document.add_heading(heading, level=1)
    _add_markdown_content(document, draft.get("content_md") or "")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _chapter_filename(draft: dict, index: int, suffix: str) -> str:
    chapter_code = str(draft.get("chapter_code") or "").strip()
    chapter_title = str(draft.get("chapter_title") or "").strip()
    code_segment = _safe_filename_segment(chapter_code, fallback=f"ch{index + 1:02d}")
    title_segment = _safe_filename_segment(chapter_title, fallback="")
    base = f"{index + 1:02d}_{code_segment}"
    if title_segment:
        base = f"{base}_{title_segment}"
    return f"{base}{suffix}"


def _default_zip_path(project_id: UUID, suffix: str) -> Path:
    return EXPORT_ROOT / str(project_id) / f"bid-chapters-{suffix}-{project_id}.zip"


def render_chapter_docx_zip(
    conn: Connection,
    *,
    project_id: UUID,
    output_path: Path | None = None,
) -> Path:
    """Render each chapter draft into its own .docx and pack into a zip."""
    drafts = _load_chapter_drafts(conn, project_id)
    if not drafts:
        raise ValueError("no chapter drafts found for project")

    project_name = _load_project_name(conn, project_id)
    if output_path is None:
        output_path = _default_zip_path(project_id, "docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    work_dir = output_path.parent / f"_chapters-docx-{project_id}"
    work_dir.mkdir(parents=True, exist_ok=True)

    chapter_paths: list[tuple[str, Path]] = []
    for index, draft in enumerate(drafts):
        filename = _chapter_filename(draft, index, ".docx")
        chapter_path = work_dir / filename
        _render_chapter_docx(project_name=project_name, draft=draft, output_path=chapter_path)
        chapter_paths.append((filename, chapter_path))

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for filename, path in chapter_paths:
            archive.write(path, filename)

    logger.info(
        "chapter_docx_zip_rendered",
        project_id=str(project_id),
        output=str(output_path),
        chapter_count=len(chapter_paths),
    )
    return output_path


def render_chapter_doc_zip(
    conn: Connection,
    *,
    project_id: UUID,
    output_path: Path | None = None,
) -> Path:
    """Render each chapter draft into a .docx, convert to legacy .doc, pack into zip."""
    drafts = _load_chapter_drafts(conn, project_id)
    if not drafts:
        raise ValueError("no chapter drafts found for project")

    project_name = _load_project_name(conn, project_id)
    if output_path is None:
        output_path = _default_zip_path(project_id, "doc")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    work_dir = output_path.parent / f"_chapters-doc-{project_id}"
    work_dir.mkdir(parents=True, exist_ok=True)

    converted: list[tuple[str, Path]] = []
    failures: list[str] = []
    for index, draft in enumerate(drafts):
        docx_filename = _chapter_filename(draft, index, ".docx")
        docx_path = work_dir / docx_filename
        _render_chapter_docx(project_name=project_name, draft=draft, output_path=docx_path)
        doc_path = convert_docx_to_doc(docx_path)
        if doc_path is None:
            failures.append(docx_filename)
            continue
        converted.append((_chapter_filename(draft, index, ".doc"), doc_path))

    if failures:
        raise RuntimeError(
            "DOC conversion unavailable or failed for chapters: " + ", ".join(failures)
        )

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for filename, path in converted:
            archive.write(path, filename)

    logger.info(
        "chapter_doc_zip_rendered",
        project_id=str(project_id),
        output=str(output_path),
        chapter_count=len(converted),
    )
    return output_path




def inspect_rendered_docx_evidence(path: Path) -> dict:
    """Inspect a rendered DOCX for page-count and chart-placeholder evidence."""

    xml_text = ""
    media_count = 0
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml_text += archive.read(name).decode("utf-8", errors="ignore")
            elif name.startswith("word/media/"):
                media_count += 1
    residual = sorted(set(CHART_PLACEHOLDER_RE.findall(xml_text)))
    return {
        "path": str(path),
        "media_count": media_count,
        "residual_chart_placeholders": residual,
        "residual_chart_placeholder_count": len(residual),
        "page_count": count_docx_pages(path),
    }


def render_export(
    conn: Connection,
    *,
    project_id: UUID,
    mode: str = EXPORT_MODE_SINGLE_DOCX,
    template_name: str = "default_technical_bid.docx",
    output_path: Path | None = None,
) -> Path:
    """Dispatch export rendering based on the requested mode."""
    if mode == EXPORT_MODE_SINGLE_DOCX:
        return render_docx(
            conn,
            project_id=project_id,
            template_name=template_name,
            output_path=output_path,
        )
    if mode == EXPORT_MODE_MULTI_DOCX_ZIP:
        return render_chapter_docx_zip(conn, project_id=project_id, output_path=output_path)
    if mode == EXPORT_MODE_MULTI_DOC_ZIP:
        return render_chapter_doc_zip(conn, project_id=project_id, output_path=output_path)
    raise ValueError(f"unsupported export mode: {mode}")
