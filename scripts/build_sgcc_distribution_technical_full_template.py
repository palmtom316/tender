#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = ROOT / "backend"
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from tender_backend.services.bid_outline_templates import SGCC_DISTRIBUTION_TECHNICAL_CHAPTERS

OUT_DIR = ROOT / "docs/samples/template_import_ready/sgcc_distribution_technical_full_package"
OUT_DOCX = OUT_DIR / "国网配网工程技术标_完整模板.docx"
REVIEW = ROOT / "docs/outputs/国网配网工程技术标模板_完整评审版.docx"
MANIFEST = OUT_DIR / "manifest.json"
README = OUT_DIR / "README.md"

ALL = [(str(ch["chapter_code"]), str(ch["chapter_title"])) for ch in SGCC_DISTRIBUTION_TECHNICAL_CHAPTERS]
TOP = [(code, title) for code, title in ALL if "." not in code]
CH8 = [(code, title) for code, title in ALL if code.startswith("8.")]
CH9 = [(code, title) for code, title in ALL if code.startswith("9.")]


def setup() -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.sections[0].header.paragraphs[0].text = ""
    return doc


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def cover(doc: Document, code: str, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(260)
    r = p.add_run(f"{code}. {title}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("（本页不编辑正文）")
    r2.font.size = Pt(10.5)
    r2.font.name = "宋体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    page_break(doc)


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def hc(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def table(doc: Document, headers: list[str], row: list[str]):
    t = doc.add_table(rows=2, cols=len(headers))
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        t.cell(0, i).text = value
    for i, value in enumerate(row):
        t.cell(1, i).text = value
    return t


def add_directory(doc: Document) -> None:
    hc(doc, "标书目录", 2)
    for code, title in ALL:
        if code == "0" or code.startswith("0."):
            continue
        indent = "    " if "." in code else ""
        p(doc, f"{indent}{code}. {title}..........第 页")


def body(doc: Document, code: str, title: str) -> None:
    if code == "0":
        hc(doc, "技术标重点资料索引", 1)
        hc(doc, "技术评分标准支撑材料", 2)
        table(doc, ["序号", "评分项名称", "资料名称", "技术标资料位置", "对应评分要求", "状态"], ["{{ row.index }}", "{{ scoring_item.name }}", "{{ material.name }}", "{{ material.location }}", "{{ scoring_item.requirement }}", "{{ row.status }}"])
        hc(doc, "技术规范书规定应该提交的材料", 2)
        table(doc, ["序号", "规范条款/来源", "资料名称", "技术标资料位置", "对应规范要求", "状态"], ["{{ row.index }}", "{{ spec.source }}", "{{ material.name }}", "{{ material.location }}", "{{ spec.requirement }}", "{{ row.status }}"])
        page_break(doc)
        add_directory(doc)
    elif code == "1":
        h(doc, "技术偏差表", 1)
        table(doc, ["序号", "偏差事项", "招标文件要求", "投标文件响应", "偏差说明"], ["1", "无偏差事项", "招标文件全部要求", "投标文件全部响应", "无偏差"])
        table(doc, ["序号", "偏差事项", "招标文件要求", "投标文件响应", "偏差说明"], ["", "以下无正文", "", "", ""])
        table(doc, ["序号", "偏差事项", "招标文件要求", "投标文件响应", "偏差说明"], ["", "其他", "/", "/", "/"])
        p(doc, "投标人声明：针对本招标标的，除本表已列明偏差外，我们接受招标文件规定的其余全部技术条件，并承诺按照招标文件规定的技术条件提供对应服务。")
    elif code == "2":
        h(doc, "关于施工项目人员执业合规的承诺函", 1)
        p(doc, "我公司承诺：")
        p(doc, "本项目施工阶段所配备的项目负责人的执业符合《注册建造师管理规定》要求，且不存在同时担任两个及两个以上建设工程项目的项目负责人的情形。")
    elif code == "3":
        h(doc, "工期/服务期承诺函", 1)
        p(doc, "我公司承诺本项目工期/服务期响应、进度计划满足招标文件要求。")
        p(doc, "{{ tender.construction_period }}")
    elif code == "4":
        h(doc, "资质情况", 1)
        p(doc, "提供符合第一章招标公告——投标人专用资格要求中“资质要求”的证明文件，如企业资质等级证书影印件或官方网站查询结果等。")
        p(doc, "{{ asset:qualification_assets:n }}")
    elif code == "5":
        h(doc, "业绩情况", 1)
        p(doc, "{{ selected_performance_summary }}")
        h(doc, "5.1. 类似工程业绩情况汇总表", 2)
        table(doc, ["序号", "工程名称", "业主单位", "合同签订日期", "完成/在建状态", "证明材料位置"], ["{{ row.index }}", "{{ performance.project_name }}", "{{ performance.client_name }}", "{{ performance.contract_date }}", "{{ performance.status }}", "{{ asset:performance_assets:n }}"])
        h(doc, "5.2. 近年完成的类似项目情况及证明材料", 2)
        table(doc, ["项目名称", "工程所在地", "业主名称", "合同金额", "开竣工日期", "承担工作", "证明材料"], ["{{ performance.project_name }}", "{{ performance.location }}", "{{ performance.client_name }}", "{{ performance.contract_amount }}", "{{ performance.period }}", "{{ performance.scope }}", "{{ asset:performance_completed_assets:n }}"])
        h(doc, "5.3. 正在施工的和新承接的类似项目情况及证明材料", 2)
        table(doc, ["项目名称", "工程所在地", "业主名称", "合同金额", "合同签订日期", "当前状态", "证明材料"], ["{{ performance.project_name }}", "{{ performance.location }}", "{{ performance.client_name }}", "{{ performance.contract_amount }}", "{{ performance.contract_date }}", "{{ performance.status }}", "{{ asset:performance_ongoing_assets:n }}"])
    elif code == "6":
        h(doc, "现场管理机构设置", 1)
        p(doc, "项目管理机构附表（不限于）：")
        h(doc, "表1 项目管理机构组成表", 2)
        table(doc, ["层级", "管理专业", "职务", "姓名", "职称", "执业或职业资格证明", "备注"], ["{{ personnel.level }}", "{{ personnel.management_specialty }}", "{{ personnel.role }}", "{{ person.full_name }}", "{{ person.title }}", "{{ person.qualification }}", "{{ personnel.note }}"])
        p(doc, "{{ personnel_table }}")
        p(doc, "{{ asset:personnel_certificates:n }}")
    elif code == "7":
        h(doc, "其他资格条件情况", 1)
        p(doc, "提供符合第一章招标公告——投标人专用资格要求中“其他”处要求的相关证明文件。")
        p(doc, "{{ other_qualification_requirements }}")
        p(doc, "{{ asset:other_qualification_assets:n }}")
    elif code == "8":
        h(doc, "施工方案与技术措施", 1)
        p(doc, "{{ ai:chapter_8_longform }}")
        for child_code, child_title in CH8:
            h(doc, f"{child_code}. {child_title}", 2)
            p(doc, f"{{{{ ai:{child_code} }}}}")
        for chart in ["construction_flow", "risk_matrix", "quality_system", "safety_system", "schedule_gantt", "equipment_table", "closure_flow"]:
            p(doc, f"{{{{ chart:{chart} }}}}")
    elif code == "9":
        h(doc, "工作规划描述", 1)
        p(doc, "{{ ai:chapter_9_work_plan }}")
        for child_code, child_title in CH9:
            h(doc, f"{child_code}. {child_title}", 2)
            p(doc, f"{{{{ ai:{child_code} }}}}")
        for chart in ["responsibility_matrix", "risk_matrix", "response_matrix", "interface_table", "indicator_table"]:
            p(doc, f"{{{{ chart:{chart} }}}}")
    elif code == "10":
        h(doc, "10.1. 质量保障措施", 2)
        p(doc, "{{ ai:10.1_quality_assurance }}")
        p(doc, "{{ chart:quality_system }}")
        p(doc, "{{ chart:closure_flow }}")
        h(doc, "10.2. 安全和绿色施工保障措施", 2)
        p(doc, "{{ ai:10.2_safety_green }}")
        p(doc, "{{ chart:safety_system }}")
        p(doc, "{{ chart:risk_matrix }}")
        h(doc, "10.3. 工程进度计划及保证措施", 2)
        p(doc, "{{ ai:10.3_schedule_assurance }}")
        p(doc, "{{ chart:schedule_gantt }}")
        p(doc, "{{ chart:critical_path }}")
    elif code == "11":
        h(doc, "履约评价证明材料", 1)
        table(doc, ["序号", "评价单位", "项目/服务名称", "评价结果", "评价年度", "证明材料位置"], ["{{ row.index }}", "{{ evaluation.owner_name }}", "{{ evaluation.project_name }}", "{{ evaluation.result }}", "{{ evaluation.year }}", "{{ asset:performance_evaluation_assets:n }}"])
        p(doc, "{{ performance_evaluation_rows }}")
    elif code == "12":
        h(doc, "施工外包管理", 1)
        p(doc, "投标人拟在中标后将中标项目的专业工程或劳务作业发包给其他外包商完成的活动。施工业务外包分为专业工程分包（以下简称“专业分包”）和劳务作业分包（以下简称“劳务分包”）。")
        p(doc, "外包计划管理（包括拟外包内容及方式）；如拟外包的，对外包单位、外包项目部、外包人员管理，外包安全质量控制；农民工工资支付管理。")
        h(doc, "附表：拟外包项目情况表", 2)
        table(doc, ["工程拟外包的项目、部位", "外包性质（劳务、专业）", "主要内容", "预估造价（万元）", "拟外包商资质等级说明"], ["{{ subcontract.scope }}", "{{ subcontract.mode }}", "{{ subcontract.content }}", "{{ subcontract.estimated_amount }}", "{{ subcontract.qualification_note }}"])
        h(doc, "承诺书", 2)
        p(doc, "我公司参与此次招标采购活动，若我公司有幸中选，我公司承诺按招标文件和合同要求依法合规实施施工外包管理，不转包、不违法分包、不挂靠。")
        p(doc, "投标人：{{ company.company_name }}")
    elif code == "13":
        h(doc, "其他", 1)
        p(doc, "{{ asset:other_technical_assets:n }}")


def build(path: Path) -> None:
    doc = setup()
    first = True
    for code, title in TOP:
        if not first:
            page_break(doc)
        first = False
        if code != "0":
            cover(doc, code, title)
        body(doc, code, title)
    doc.save(path)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    REVIEW.parent.mkdir(parents=True, exist_ok=True)
    build(OUT_DOCX)
    build(REVIEW)
    manifest = {
        "package_name": "国网配网工程技术标完整模板",
        "package_key": "sgcc_distribution_technical_full_v1",
        "package_type": "technical",
        "category_code": "sgcc_distribution",
        "chapter_count": len(ALL),
        "output_docx": str(OUT_DOCX.relative_to(ROOT)),
        "chapters": [
            {"chapter_code": code, "chapter_name": title, "render_mode": "single_docx_section"}
            for code, title in ALL
        ],
        "template_policy": {
            "chapter_0_is_index_directory_without_cover": True,
            "cover_page_top_level_chapters": [code for code, _title in TOP if code != "0"],
            "ai_generated_chapters": ["8", "9", "10.1", "10.2", "10.3"],
            "no_internal_instruction_text_in_docx": True,
            "historical_performance_removed": True,
            "historical_personnel_removed": True,
            "images_removed": True,
        },
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    README.write_text(
        "完整技术标模板：第0章为目录和索引章，不设置章节封面；第1-13章一级章首页为“章号. 标题 / （本页不编辑正文）”；第8/9/10章只保留AI生成锚点和图表占位符，不含历史正文。\n",
        encoding="utf-8",
    )
    print(OUT_DOCX.relative_to(ROOT))
    print(REVIEW.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
