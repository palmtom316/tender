#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = ROOT / "backend"
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from tender_backend.services.bid_outline_templates import SGCC_DISTRIBUTION_BUSINESS_CHAPTERS

BUSINESS_PACKAGE_KEY = "sgcc_distribution_business_full_v1"
OUT_DIR = ROOT / "docs/samples/template_import_ready/sgcc_distribution_business_full_package"
OUT_DOCX = OUT_DIR / "国网配网工程商务标_完整模板.docx"
REVIEW = ROOT / "docs/outputs/国网配网工程商务标模板_评审版.docx"
MANIFEST = OUT_DIR / "manifest.json"
README = OUT_DIR / "README.md"

ALL = [(str(ch["chapter_code"]), str(ch["chapter_title"]), ch.get("parent_code")) for ch in SGCC_DISTRIBUTION_BUSINESS_CHAPTERS]
TOP = [(code, title) for code, title, parent in ALL if not parent]
CHILDREN_BY_PARENT: dict[str, list[tuple[str, str]]] = {}
for code, title, parent in ALL:
    if parent:
        CHILDREN_BY_PARENT.setdefault(str(parent), []).append((code, title))


def setup() -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = Pt(10.5)
    return doc


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def cover(doc: Document, code: str, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(260)
    r = p.add_run(f"{code}. {title}")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("（本页不编辑正文）")
    r2.font.size = Pt(10.5)
    r2.font.name = "宋体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    page_break(doc)


def heading(doc: Document, code: str, title: str, level: int) -> None:
    doc.add_heading(f"{code}. {title}", level=level)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def table(doc: Document, headers: list[str], row: list[str]) -> None:
    t = doc.add_table(rows=2, cols=len(headers))
    t.style = "Table Grid"
    for i, value in enumerate(headers):
        t.cell(0, i).text = value
    for i, value in enumerate(row):
        t.cell(1, i).text = value


def _children_recursive(code: str) -> Iterable[tuple[str, str]]:
    for child_code, child_title in CHILDREN_BY_PARENT.get(code, []):
        yield child_code, child_title
        yield from _children_recursive(child_code)


def _asset_placeholder(code: str, title: str) -> str:
    if code == "3":
        return "{{ asset:business_license:1 }}"
    if code == "4":
        return "{{ asset:legal_representative_identity:1 }}"
    if code.startswith("8"):
        return "{{ asset:financial_statement:n }}"
    if code.startswith("12"):
        return "{{ asset:management_system_certificates:n }}"
    if code.startswith("13"):
        return "{{ asset:esg_environment_assets:n }}"
    if code == "14":
        return "{{ asset:green_certificate_assets:n }}"
    if code == "15":
        return "{{ asset:technology_achievement_assets:n }}"
    if code == "16":
        return "{{ asset:innovation_policy_assets:n }}"
    if code == "17":
        return "{{ asset:rd_team_assets:n }}"
    if code == "18":
        return "{{ asset:quality_award_assets:n }}"
    if code == "19":
        return "{{ asset:high_tech_enterprise_certificate:1 }}"
    if code == "20":
        return "{{ asset:name_change_assets:n }}"
    if code == "22":
        return "{{ asset:tax_rate_support_assets:n }}"
    if code.startswith("23"):
        return "{{ asset:bid_bond_assets:n }}"
    if code.startswith("24"):
        return "{{ asset:other_business_assets:n }}"
    if "信用" in title:
        return "{{ asset:credit_report:n }}"
    return "{{ asset:business_supporting_assets:n }}"


def _section_body(doc: Document, code: str, title: str) -> None:
    if code == "1":
        table(doc, ["序号", "招标文件条款", "招标文件要求", "投标文件响应", "偏差说明"], ["{{ row.index }}", "{{ clause.no }}", "{{ clause.requirement }}", "{{ response.text }}", "{{ deviation.note }}"])
        para(doc, "{{ business_deviation_rows }}")
    elif code == "2":
        para(doc, "致：{{ tender.owner_name }}")
        para(doc, "我方承诺参加本项目投标活动不存在招标文件规定的违法失信情形，所提交资料真实、合法、有效，并接受招标人核验。")
        para(doc, "投标人：{{ company.company_name }}")
    elif code == "4":
        table(doc, ["姓名", "性别", "职务", "身份证件号码", "单位名称"], ["{{ legal_rep.name }}", "{{ legal_rep.gender }}", "{{ legal_rep.position }}", "{{ legal_rep.id_no }}", "{{ company.company_name }}"])
        para(doc, _asset_placeholder(code, title))
    elif code == "5.1":
        table(doc, ["单位名称", "统一社会信用代码", "注册地址", "法定代表人", "联系人", "联系电话"], ["{{ company.company_name }}", "{{ company.credit_code }}", "{{ company.registered_address }}", "{{ legal_rep.name }}", "{{ contact.name }}", "{{ contact.phone }}"])
    elif code == "5.2":
        table(doc, ["序号", "查询事项", "查询结果", "查询时间", "证明材料位置"], ["{{ row.index }}", "{{ query.item }}", "{{ query.result }}", "{{ query.date }}", "{{ asset:safety_quality_query_assets:n }}"])
    elif code == "6.1":
        table(doc, ["序号", "姓名", "拟任岗位", "证书名称", "证书编号", "证明材料位置"], ["{{ row.index }}", "{{ person.name }}", "{{ person.role }}", "{{ certificate.name }}", "{{ certificate.no }}", "{{ asset:personnel_certificates:n }}"])
        para(doc, "{{ personnel_resume_rows }}")
    elif code == "7":
        para(doc, "我方已按招标文件要求核查与国家电网公司系统人员关系情况，相关说明和承诺如下：")
        para(doc, "{{ sgcc_personnel_relationship_statement }}")
    elif code == "9":
        para(doc, "{{ consortium_agreement_or_not_applicable }}")
    elif code == "10":
        para(doc, "{{ asset:bank_account_document:1 }}")
    elif code == "21":
        para(doc, "{{ small_taxpayer_statement_or_not_applicable }}")
    elif code == "23.1":
        table(doc, ["序号", "保证金形式", "金额", "缴纳/出具时间", "证明材料位置"], ["{{ row.index }}", "{{ bond.type }}", "{{ bond.amount }}", "{{ bond.issued_at }}", "{{ asset:bid_bond_assets:n }}"])
    else:
        para(doc, _asset_placeholder(code, title))


def body(doc: Document, code: str, title: str) -> None:
    heading(doc, code, title, 1)
    _section_body(doc, code, title)
    for child_code, child_title in _children_recursive(code):
        level = 2 if child_code.count(".") == 1 else 3
        heading(doc, child_code, child_title, level)
        _section_body(doc, child_code, child_title)


def build(path: Path) -> None:
    doc = setup()
    first = True
    for code, title in TOP:
        if not first:
            page_break(doc)
        first = False
        cover(doc, code, title)
        body(doc, code, title)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    REVIEW.parent.mkdir(parents=True, exist_ok=True)
    build(OUT_DOCX)
    build(REVIEW)
    manifest = {
        "package_name": "国网配网工程商务标完整模板",
        "package_key": BUSINESS_PACKAGE_KEY,
        "package_type": "business",
        "category_code": "sgcc_distribution",
        "chapter_count": len(ALL),
        "output_docx": str(OUT_DOCX.relative_to(ROOT)),
        "chapters": [
            {"chapter_code": code, "chapter_name": title, "render_mode": "single_docx_section"}
            for code, title, _parent in ALL
        ],
        "template_policy": {
            "cover_page_each_top_level_chapter": True,
            "single_docx_package": True,
            "no_internal_instruction_text_in_docx": True,
            "historical_business_facts_removed": True,
            "images_removed": True,
        },
    }
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    README.write_text(
        "完整商务标模板：按商务标1-24章源头目录生成；一级章首页为“章号. 标题 / （本页不编辑正文）”；正文仅保留固定承诺文字、表格结构和资料占位符，不含历史正文、图片或真实证明材料。\n",
        encoding="utf-8",
    )
    print(OUT_DOCX.relative_to(ROOT))
    print(REVIEW.relative_to(ROOT))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
