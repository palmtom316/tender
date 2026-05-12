#!/usr/bin/env python3
"""Generate SGCC technical bid chapters 8/9/10 with DeepSeek and save DOCX."""

from __future__ import annotations

import argparse
import os
import re
import sys
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openai import OpenAI

ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

from tender_backend.services.deepseek_api import (
    DEEPSEEK_BASE_URL,
    DEEPSEEK_V4_MAX_REASONING_EFFORT,
    DEEPSEEK_V4_PRO_MODEL,
    deepseek_v4_openai_sdk_options,
    is_deepseek_v4_model,
)


SAMPLES = ROOT / "docs" / "samples"

PROMPT_FILES = {
    "8": "配网施工方案及技术措施提示词.md",
    "9": "配网工作规划描述提示词.md",
    "10.1": "配网质量保证措施提示词.md",
    "10.2": "配网安全与绿色施工保障措施提示词.md",
    "10.3": "配网工程进度计划及保证措施提示词.md",
}

OUTPUT_ORDER = ("8", "9", "10.1", "10.2", "10.3")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def compact_prompt(text: str, limit: int | None = None) -> str:
    if limit is None or len(text) <= limit:
        return text
    marker = "## 六、"
    head, _, tail = text.partition(marker)
    compact = head[:9000] + "\n\n" + marker + tail[:9000]
    return compact[:limit]


def build_messages(project_name: str, section_key: str, catalog: str, prompt_text: str) -> list[dict[str, str]]:
    section_names = {
        "8": "第8章《施工方案与技术措施》",
        "9": "第9章《工作规划描述》",
        "10.1": "第10章第10.1节《质量保证措施》",
        "10.2": "第10章第10.2节《安全和绿色施工保障措施》",
        "10.3": "第10章第10.3节《工程进度计划及保证措施》",
    }
    project_context = f"""
项目名称：{project_name}
技术标目录：
{catalog}

已知项目事实：
- 招标/采购对象：国家电网重庆市区分公司2026年年度配网工程框架协议技术标。
- 采购类型：年度框架协议。
- 建设地点：重庆市区范围内，具体区县、站点、线路、工程量以单项工程任务书和发包人确认文件为准。
- 工程类型：10kV及以下配网工程相关施工、改造、检修配套及框架协议内派工项目。
- 暗标控制：按暗标口径处理，不出现投标人名称、人员姓名、联系人、地址、企业标识和无法佐证的特有信息。
- 未提供招标编号、标段号、具体工程量、明确工期、人员设备清单、标准条款版本、停电窗口和外部许可；涉及这些内容时写“按招标文件、任务书和发包人确认要求执行”或列入待补充，不得编造。
- 不出现价格、报价、最高限价、单价、总价、产值等商务价格内容。
""".strip()
    user_prompt = f"""
请生成{section_names[section_key]}的正式技术标正文。

{project_context}

本节/本章写作提示词如下，必须遵守其中的章节层级、边界、禁用内容和自检要求：
{compact_prompt(prompt_text)}

输出要求：
1. 只输出可直接放入投标文件的 Markdown 正文，不写说明、寒暄或模型自述。
2. 使用中文正式投标文件语气。
3. 表格使用 Markdown 表格；必要时使用清单。
4. 内容要尽量完整，但不得编造没有来源的数量、日期、人员、设备型号、证书编号、标准版本、审批时限、许可结果和量化承诺。
5. 保持章节编号准确；第10章三个子项分别输出 10.1、10.2、10.3 层级，不要重复输出整个第10章标题以外的无关章节。
""".strip()
    return [
        {
            "role": "system",
            "content": "你是国网配网工程技术标编写专家，输出必须可追溯、合规、适合DOCX导出。",
        },
        {"role": "user", "content": user_prompt},
    ]


def call_deepseek(
    client: OpenAI,
    model: str,
    messages: list[dict[str, str]],
    max_tokens: int,
    *,
    reasoning_effort: str | None = None,
    thinking_enabled: bool | None = None,
) -> tuple[str, object]:
    extra_options = {}
    if is_deepseek_v4_model(model):
        extra_options = deepseek_v4_openai_sdk_options(
            thinking_enabled=thinking_enabled,
            reasoning_effort=reasoning_effort,
        )
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        max_tokens=max_tokens,
        **extra_options,
    )
    content = response.choices[0].message.content or ""
    return content.strip(), response.usage


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("第 ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    paragraph.add_run(" 页")


def apply_doc_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = document.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for section in document.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        section.header.paragraphs[0].text = "技术标"
        add_page_number(section.footer.paragraphs[0])


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    value = line.strip()
    return bool(value) and set(value.replace("|", "").replace(":", "").replace("-", "").strip()) == set()


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines if not is_table_separator(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = row[col_idx] if col_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.font.bold = True


def add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_lines)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
        elif re.match(r"^[-*]\s+", line):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)
        index += 1


def write_docx(project_name: str, outputs: dict[str, str], output_path: Path) -> None:
    document = Document()
    apply_doc_style(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(project_name)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(18)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("技术标第8、9、10章")
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16)
    document.add_page_break()

    for idx, key in enumerate(OUTPUT_ORDER):
        if idx:
            document.add_page_break()
        add_markdown(document, outputs[key])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--project-name", required=True)
    parser.add_argument("--model", default=os.environ.get("DEEPSEEK_MODEL", DEEPSEEK_V4_PRO_MODEL))
    parser.add_argument("--base-url", default=os.environ.get("DEEPSEEK_BASE_URL", DEEPSEEK_BASE_URL))
    parser.add_argument("--output", required=True)
    parser.add_argument("--max-tokens", type=int, default=64000)
    parser.add_argument("--reasoning-effort", default=os.environ.get("DEEPSEEK_REASONING_EFFORT", DEEPSEEK_V4_MAX_REASONING_EFFORT))
    parser.add_argument("--disable-thinking", action="store_true")
    parser.add_argument("--sleep", type=float, default=1.0)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    api_key = os.environ.get("DEEPSEEK_API_KEY")
    if not api_key:
        print("DEEPSEEK_API_KEY is required", file=sys.stderr)
        return 2

    catalog = read_text(SAMPLES / "国网公司配网工程技术标目录.md")
    prompts = {key: read_text(SAMPLES / filename) for key, filename in PROMPT_FILES.items()}
    client = OpenAI(api_key=api_key, base_url=args.base_url, timeout=600, max_retries=1)

    outputs: dict[str, str] = {}
    usage_summary: dict[str, dict[str, int]] = {}
    for key in OUTPUT_ORDER:
        print(f"Generating {key} with {args.model}...", flush=True)
        content, usage = call_deepseek(
            client,
            args.model,
            build_messages(args.project_name, key, catalog, prompts[key]),
            args.max_tokens,
            reasoning_effort=args.reasoning_effort,
            thinking_enabled=False if args.disable_thinking else True,
        )
        outputs[key] = content
        usage_summary[key] = {
            "prompt_tokens": int(getattr(usage, "prompt_tokens", 0) or 0),
            "completion_tokens": int(getattr(usage, "completion_tokens", 0) or 0),
            "total_tokens": int(getattr(usage, "total_tokens", 0) or 0),
        }
        print(
            f"Finished {key}: {len(content)} chars, "
            f"{usage_summary[key]['total_tokens']} tokens",
            flush=True,
        )
        time.sleep(args.sleep)

    output_path = Path(args.output)
    write_docx(args.project_name, outputs, output_path)
    md_path = output_path.with_suffix(".md")
    md_path.write_text("\n\n---\n\n".join(outputs[key] for key in OUTPUT_ORDER), encoding="utf-8")
    print(f"DOCX: {output_path}")
    print(f"Markdown backup: {md_path}")
    print(f"Usage: {usage_summary}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
