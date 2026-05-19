#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document


_FILENAME_RE = re.compile(r"^(?P<code>\d+(?:\.\d+)*)(?:\.|．)(?P<title>.+?)\.docx$", re.IGNORECASE)
_PARAGRAPH_HEADING_RE = re.compile(r"^\s*(?P<code>\d+(?:\.\d+)*)(?:\.|．)(?P<title>[^\d.．].+?)\s*$")
_DUPLICATED_TOP_LEVEL_CODES = {"9.9": "9"}


@dataclass(frozen=True)
class PackageItem:
    source_path: Path
    original_code: str
    normalized_code: str
    title: str


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _code_sort_key(code: str) -> tuple[int, ...]:
    return tuple(int(part) for part in code.split("."))


def _parse_item(path: Path) -> PackageItem | None:
    match = _FILENAME_RE.match(path.name)
    if match is None:
        return None
    original_code = match.group("code")
    normalized_code = _DUPLICATED_TOP_LEVEL_CODES.get(original_code, original_code)
    return PackageItem(
        source_path=path,
        original_code=original_code,
        normalized_code=normalized_code,
        title=match.group("title").strip(),
    )


def list_package_items(source_dir: str | Path) -> list[PackageItem]:
    root = Path(source_dir).expanduser().resolve()
    if not root.exists():
        raise FileNotFoundError(f"source directory not found: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"source path is not a directory: {root}")

    items = [_parse_item(path) for path in root.glob("*.docx")]
    parsed = [item for item in items if item is not None]
    if not parsed:
        raise ValueError(f"no numbered DOCX files found in: {root}")
    return sorted(parsed, key=lambda item: (_code_sort_key(item.normalized_code), item.source_path.name))


def _paragraph_heading_code(text: str) -> str | None:
    match = _PARAGRAPH_HEADING_RE.match(text or "")
    if match is None:
        return None
    return match.group("code")


def _looks_like_source_heading(text: str, item: PackageItem) -> bool:
    stripped = (text or "").strip()
    prefixes = (
        f"{item.original_code}.",
        f"{item.original_code}．",
        f"{item.normalized_code}.",
        f"{item.normalized_code}．",
    )
    return stripped.startswith(prefixes)


def _body_elements_without_leading_heading(source: Document, item: PackageItem) -> list:
    elements = list(source.element.body)
    while elements:
        element = elements[0]
        if element.tag.endswith("}sectPr"):
            elements.pop(0)
            continue
        text = "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t")).strip()
        code = _paragraph_heading_code(text)
        if code in {item.original_code, item.normalized_code} or _looks_like_source_heading(text, item):
            elements.pop(0)
            continue
        break
    return [element for element in elements if not element.tag.endswith("}sectPr")]


def merge_package(source_dir: str | Path, output_docx: str | Path) -> dict[str, Any]:
    items = list_package_items(source_dir)
    output_path = Path(output_docx).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    merged = Document()
    if merged.paragraphs and not merged.paragraphs[0].text:
        paragraph = merged.paragraphs[0]
        paragraph._element.getparent().remove(paragraph._element)

    evidence_items: list[dict[str, Any]] = []
    for index, item in enumerate(items):
        if index:
            merged.add_page_break()
        merged.add_heading(f"{item.normalized_code}. {item.title}", level=1)
        source = Document(str(item.source_path))
        for element in _body_elements_without_leading_heading(source, item):
            merged.element.body.append(deepcopy(element))
        evidence_items.append(
            {
                "source_filename": item.source_path.name,
                "source_sha256": _sha256(item.source_path),
                "source_size_bytes": item.source_path.stat().st_size,
                "original_code": item.original_code,
                "normalized_code": item.normalized_code,
                "renormalized": item.original_code != item.normalized_code,
            }
        )

    merged.save(str(output_path))
    return {
        "source_dir": str(Path(source_dir).expanduser().resolve()),
        "output_docx": str(output_path),
        "output_sha256": _sha256(output_path),
        "output_size_bytes": output_path.stat().st_size,
        "item_count": len(items),
        "merge_order_codes": [item.normalized_code for item in items],
        "items": evidence_items,
        "merge_note": "content text is intentionally omitted from evidence",
    }


def to_json_text(evidence: dict[str, Any]) -> str:
    return json.dumps(evidence, ensure_ascii=False, indent=2, default=str) + "\n"


def main() -> int:
    parser = argparse.ArgumentParser(description="Merge a confidential multi-DOCX business package into one local DOCX.")
    parser.add_argument("--source-dir", required=True, help="Directory containing numbered business DOCX files")
    parser.add_argument("--output-docx", required=True, help="Merged confidential DOCX output path")
    parser.add_argument("--evidence-output", required=True, help="Sanitized merge evidence JSON output path")
    args = parser.parse_args()

    evidence = merge_package(args.source_dir, args.output_docx)
    evidence_output = Path(args.evidence_output).expanduser().resolve()
    evidence_output.parent.mkdir(parents=True, exist_ok=True)
    evidence_output.write_text(to_json_text(evidence), encoding="utf-8")
    print(f"Wrote merged DOCX to {evidence['output_docx']}")
    print(f"Wrote sanitized merge evidence to {evidence_output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
